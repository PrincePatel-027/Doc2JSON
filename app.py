"""
Doc2JSON v2 -- Advanced Document to JSON Converter for AI Training Data
Handles entire books with multi-pass extraction, Mistral AI OCR for
images/diagrams, chapter detection, header/footer removal, and quality validation.
"""

import os
import json
import re
import hashlib
import uuid
import base64
import io
import math
import logging
import copy
from datetime import datetime, timezone
from pathlib import Path
from collections import Counter
import concurrent.futures

from flask import Flask, request, jsonify, render_template, send_file
from PIL import Image, ImageEnhance, ImageFilter

# PDF engines
import fitz  # PyMuPDF -- primary engine
import pdfplumber  # secondary engine -- better for tables

# ── Mistral AI OCR (primary OCR engine) ──────────────────────────────
MISTRAL_AVAILABLE = False
try:
    from mistralai.client import Mistral
    MISTRAL_AVAILABLE = True
except Exception as _me:
    print(f'  [WARN] Mistral import failed: {_me}')

# ── Tesseract OCR (fallback) ─────────────────────────────────────────
TESSERACT_AVAILABLE = False
try:
    import pytesseract
    pytesseract.get_tesseract_version()
    TESSERACT_AVAILABLE = True
except Exception:
    pass

# Combined flag
OCR_AVAILABLE = MISTRAL_AVAILABLE or TESSERACT_AVAILABLE

# ── AI Validation Support ───────────────────────────────────────────
OPENAI_AVAILABLE = False
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except Exception as _oe:
    print(f'  [WARN] OpenAI import failed: {_oe}')

GROQ_BASE_URL = 'https://api.groq.com/openai/v1'
GROQ_MODELS = [
    'llama-3.3-70b-versatile',
    'llama-3.1-8b-instant',
]

SARVAM_BASE_URL = 'https://api.sarvam.ai/v1/'
SARVAM_MODELS = [
    'sarvam-30b',
]

VALIDATOR_SYSTEM_PROMPT = (
    'You are a strict JSON extraction validator. You evaluate converted document JSON for '\
    'structure quality, text accuracy, and reliability. You MUST return only a single JSON '\
    'object with these keys: score, structure_score, accuracy_score, reliability_score, '\
    'summary, issues, suggestions. The scores are from 1 to 10. issues is an array of short '\
    'strings. suggestions is an array where each item has: id, title, reason, '\
    'expected_impact, operations. operations is an array of patch operations that can be '\
    'applied later. Allowed operation shapes: '\
    '{"op":"set","path":"a.b[0].c","value":<json>} '\
    '{"op":"replace_substring","path":"a.b","search":"old","replace":"new"} '\
    '{"op":"append_list","path":"a.items","value":<json>} '\
    '{"op":"remove","path":"a.b[0]"}. '\
    'Return at most 8 suggestions and at most 3 operations per suggestion. '\
    'Do not add markdown fences or any extra text.'
)

MUTUAL_DECISION_NOTE = (
    'When multiple validator providers are used, compute a consensus result by combining '\
    'their scores and prioritizing overlapping issues and compatible suggestions.'
)

# DOCX support
DOCX_AVAILABLE = False
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    pass

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s [%(levelname)s] %(message)s')
log = logging.getLogger('doc2json')

# =====================================================================
# FLASK APP
# =====================================================================

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # 200 MB for large books
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join(os.path.dirname(__file__), 'output')
app.config['IMAGE_FOLDER'] = os.path.join(os.path.dirname(__file__), 'output', 'images')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
os.makedirs(app.config['IMAGE_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf', 'txt', 'md', 'docx', 'csv', 'html'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# =====================================================================
# DOCUMENT PROCESSOR -- Multi-pass extraction engine
# =====================================================================

class DocumentProcessor:
    """
    Comprehensive document processor with:
    - Multi-pass PDF extraction (PyMuPDF + pdfplumber)
    - Mistral AI OCR for images and diagrams (world-class accuracy)
    - Quality scoring and best-text selection per page
    - Chapter and section detection
    - Header / footer / page-number removal
    - Text cleaning and validation
    """

    # Image size thresholds
    MIN_IMAGE_DIM = 60        # skip tiny icons/bullets
    MAX_IMAGE_PIXELS = 4096   # resize oversized images
    MAX_IMAGES = 500          # cap for very image-heavy docs

    # Chapter detection patterns
    CHAPTER_PATTERNS = [
        re.compile(r'^(chapter|ch\.?)\s+(\d+|[ivxlc]+)', re.IGNORECASE),
        re.compile(r'^(part)\s+(\d+|[ivxlc]+|one|two|three|four|five)', re.IGNORECASE),
        re.compile(r'^(section)\s+(\d+)', re.IGNORECASE),
        re.compile(r'^(\d+)\.\s+[A-Z]'),      # "1. Introduction"
        re.compile(r'^(\d+)\s+[A-Z][a-z]+'),   # "1 Introduction"
    ]

    def __init__(self, filepath, filename, options=None):
        self.filepath = filepath
        self.filename = filename
        self.options = options or {}
        self.ext = Path(filename).suffix.lower()

        self.pages = []
        self.images = []
        self.metadata = {}
        self.structure = {'chapters': [], 'sections': []}
        self.quality_report = {}

        # Mistral AI client (initialized if API key is provided)
        self.mistral_client = None
        api_key = self.options.get('mistral_api_key', '') or os.environ.get('MISTRAL_API_KEY', '')
        if api_key and MISTRAL_AVAILABLE:
            self.mistral_client = Mistral(api_key=api_key)
            log.info('Mistral AI OCR client initialized')

    # -----------------------------------------------------------------
    # MAIN PIPELINE
    # -----------------------------------------------------------------

    def process(self):
        """Run the full extraction pipeline and return structured data."""
        log.info(f"Processing '{self.filename}' ({self.ext})")

        # Step 1: Extract raw text
        if self.ext == '.pdf':
            self._extract_pdf_multipass()
        elif self.ext in ('.txt', '.md', '.csv', '.html'):
            self._extract_text_file()
        elif self.ext == '.docx':
            self._extract_docx()
        else:
            raise ValueError(f'Unsupported file type: {self.ext}')

        # Step 2: Clean text on every page
        self._clean_all_pages()

        # Step 3: Remove repeated headers / footers / page numbers
        if len(self.pages) >= 4:
            self._strip_headers_footers()

        # Step 4: Detect document structure (chapters, sections)
        self._detect_structure()

        # Step 5: Final validation pass
        self._validate()

        # Step 6: AI Validation pass (if configured)
        if hasattr(self, '_ai_validate'):
            self._ai_validate()

        log.info(f"Extraction complete: {len(self.pages)} pages, "
                 f"{len(self.images)} images, "
                 f"quality={self.quality_report.get('overall_score', 'N/A')}")

        return {
            'pages': self.pages,
            'images': self.images,
            'metadata': self.metadata,
            'structure': self.structure,
            'quality': self.quality_report,
        }

    # -----------------------------------------------------------------
    # PDF -- MULTI-PASS EXTRACTION
    # -----------------------------------------------------------------

    def _extract_pdf_multipass(self):
        """Extract PDF with multiple engines and merge the best text per page.
        If Mistral API key is provided, also runs Mistral OCR as a third pass."""
        log.info("Pass 1/3: PyMuPDF extraction...")
        pass_mu = self._pymupdf_extract()

        log.info("Pass 2/3: pdfplumber extraction...")
        pass_plumber = self._pdfplumber_extract()

        log.info("Pass 3/3: Merging & scoring pages...")
        self.pages = self._merge_passes(pass_mu, pass_plumber)

        # Mistral OCR pass -- process entire PDF via Mistral API
        if self.mistral_client and self.ext == '.pdf':
            log.info("Mistral AI OCR pass: processing full document...")
            self._mistral_ocr_full_pdf()

        # Image extraction (separate pass for isolation)
        if self.options.get('extract_images', True) and self.ext == '.pdf':
            log.info("Extracting images and diagrams...")
            self._extract_images()

    # -- PyMuPDF pass --------------------------------------------------

    def _pymupdf_extract(self):
        pages = []
        with fitz.open(self.filepath) as doc:
            self.metadata['total_pages'] = len(doc)
            raw_meta = doc.metadata or {}
            self.metadata['pdf_metadata'] = {
                k: str(v) for k, v in raw_meta.items() if v
            }
            for page in doc:
                text = page.get_text('text') or ''
                pages.append({
                    'page_number': page.number + 1,
                    'text': text.strip(),
                    'char_count': len(text.strip()),
                })
        return pages

    # -- pdfplumber pass -----------------------------------------------

    def _pdfplumber_extract(self):
        pages = []
        try:
            with pdfplumber.open(self.filepath) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = ''
                    try:
                        text = page.extract_text() or ''
                    except Exception:
                        pass
                    tables = []
                    try:
                        tables = page.extract_tables() or []
                    except Exception:
                        pass
                    pages.append({
                        'page_number': i + 1,
                        'text': text.strip(),
                        'char_count': len(text.strip()),
                        'tables': tables,
                    })
        except Exception as e:
            log.warning(f"pdfplumber failed: {e}")
        return pages

    # -- Merge passes --------------------------------------------------

    def _merge_passes(self, pass1, pass2):
        """Per-page: pick the extraction with the higher quality score,
        but always keep tables from pdfplumber."""
        merged = []
        n = max(len(pass1), len(pass2))
        scores_1, scores_2 = [], []

        for i in range(n):
            p1 = pass1[i] if i < len(pass1) else None
            p2 = pass2[i] if i < len(pass2) else None

            s1 = self._quality_score(p1['text']) if p1 else 0
            s2 = self._quality_score(p2['text']) if p2 else 0
            scores_1.append(s1)
            scores_2.append(s2)

            best = dict(p1) if s1 >= s2 and p1 else dict(p2) if p2 else dict(p1)

            # Always keep tables from pdfplumber
            if p2 and p2.get('tables'):
                best['tables'] = p2['tables']
            elif 'tables' not in best:
                best['tables'] = []

            best['quality_scores'] = {
                'pymupdf': round(s1, 4),
                'pdfplumber': round(s2, 4),
                'selected_engine': 'pymupdf' if s1 >= s2 else 'pdfplumber',
            }
            merged.append(best)

        # Overall quality summary
        if scores_1:
            self.quality_report['avg_score_pymupdf'] = round(
                sum(scores_1) / len(scores_1), 4)
        if scores_2:
            self.quality_report['avg_score_pdfplumber'] = round(
                sum(scores_2) / len(scores_2), 4)

        return merged

    # -----------------------------------------------------------------
    # MISTRAL AI OCR -- Full PDF + Individual Images
    # -----------------------------------------------------------------

    def _mistral_ocr_full_pdf(self):
        """Process the entire PDF through Mistral's OCR API.
        Enriches page text and extracts images with OCR in one shot."""
        try:
            # Step 1: Upload the file to Mistral
            with open(self.filepath, 'rb') as f:
                uploaded = self.mistral_client.files.upload(
                    file={'file_name': self.filename, 'content': f.read()},
                    purpose='ocr',
                )
            log.info(f"Uploaded to Mistral (file_id={uploaded.id})")

            # Step 2: Get a signed URL
            signed = self.mistral_client.files.get_signed_url(
                file_id=uploaded.id)

            # Step 3: Run OCR
            ocr_resp = self.mistral_client.ocr.process(
                model='mistral-ocr-latest',
                document={
                    'type': 'document_url',
                    'document_url': signed.url,
                },
                include_image_base64=True,
            )

            # Step 4: Merge Mistral's results into our pages
            mistral_pages_used = 0
            for mpage in ocr_resp.pages:
                idx = mpage.index  # 0-based
                md_text = mpage.markdown or ''

                if idx < len(self.pages):
                    existing = self.pages[idx]
                    ms = self._quality_score(md_text)
                    es = self._quality_score(existing['text'])

                    # Use Mistral text if it's significantly better quality (>0.1 difference)
                    # This prevents truncated results with high 'per-word' quality
                    # from replacing slightly lower quality but more complete text.
                    if ms > es + 0.1:
                        existing['text'] = md_text
                        existing['char_count'] = len(md_text)
                        if 'quality_scores' not in existing:
                            existing['quality_scores'] = {}
                        existing['quality_scores']['mistral'] = round(ms, 4)
                        existing['quality_scores']['selected_engine'] = 'mistral'
                        mistral_pages_used += 1
                    else:
                        if 'quality_scores' not in existing:
                            existing['quality_scores'] = {}
                        existing['quality_scores']['mistral'] = round(ms, 4)

                # Collect any images Mistral found on this page
                if hasattr(mpage, 'images') and mpage.images:
                    for mi_idx, mimg in enumerate(mpage.images):
                        img_id = getattr(mimg, 'id', f'mistral_p{idx}_img{mi_idx}')
                        img_b64 = getattr(mimg, 'image_base64', '') or ''

                        if not img_b64:
                            continue

                        # Decode to get dimensions
                        try:
                            img_bytes = base64.b64decode(img_b64)
                            pil_img = Image.open(io.BytesIO(img_bytes))
                            w, h = pil_img.size
                        except Exception:
                            w, h = 0, 0
                            img_bytes = b''

                        img_filename = f"{img_id}.png"
                        if img_bytes:
                            img_path = os.path.join(
                                app.config['IMAGE_FOLDER'], img_filename)
                            with open(img_path, 'wb') as f:
                                f.write(img_bytes)

                        self.images.append({
                            'id': img_id,
                            'page': idx + 1,
                            'index_on_page': mi_idx,
                            'format': 'png',
                            'width': w,
                            'height': h,
                            'size_bytes': len(img_bytes),
                            'position': None,
                            'caption': '',
                            'overlay_text': '',
                            'ocr_text': f'[Extracted by Mistral OCR from {img_id}]',
                            'ocr_details': [{'engine': 'mistral-ocr-latest'}],
                            'image_file': img_filename,
                            'base64': img_b64,
                        })

            log.info(f"Mistral OCR: used for {mistral_pages_used}/{len(self.pages)} pages, "
                     f"found {sum(1 for img in self.images if 'mistral' in img.get('id', ''))} images")

            # Cleanup: delete the uploaded file from Mistral
            try:
                self.mistral_client.files.delete(file_id=uploaded.id)
            except Exception:
                pass

        except Exception as e:
            log.warning(f"Mistral OCR failed (falling back to local engines): {e}")

    def _mistral_ocr_image(self, pil_img, img_id='image'):
        """OCR a single image using Mistral's vision API.
        Returns (text, details_list)."""
        if not self.mistral_client:
            return '', []

        try:
            # Convert to base64
            buf = io.BytesIO()
            fmt = 'PNG' if pil_img.mode == 'RGBA' else 'JPEG'
            pil_img.save(buf, format=fmt, quality=90)
            b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
            mime = 'image/png' if fmt == 'PNG' else 'image/jpeg'

            ocr_resp = self.mistral_client.ocr.process(
                model='mistral-ocr-latest',
                document={
                    'type': 'image_url',
                    'image_url': f'data:{mime};base64,{b64}',
                },
            )

            # Collect text from all pages (usually just 1 for an image)
            texts = []
            for pg in ocr_resp.pages:
                if pg.markdown:
                    texts.append(pg.markdown)
            combined = '\n'.join(texts).strip()

            details = [{
                'engine': 'mistral-ocr-latest',
                'text': combined,
                'length': len(combined),
            }]
            return combined, details

        except Exception as e:
            log.warning(f"Mistral image OCR failed for {img_id}: {e}")
            return '', []

    # -----------------------------------------------------------------
    # IMAGE & DIAGRAM EXTRACTION
    # -----------------------------------------------------------------

    def _extract_images(self):
        """Extract images and diagrams from the PDF.
        Runs OCR on each if Tesseract is available."""
        extracted = 0
        doc_id = Path(self.filename).stem

        with fitz.open(self.filepath) as doc:
            for page in doc:
                if extracted >= self.MAX_IMAGES:
                    break

                image_list = page.get_images(full=True)
                for img_idx, img_info in enumerate(image_list):
                    if extracted >= self.MAX_IMAGES:
                        break
                    try:
                        xref = img_info[0]
                        base_img = doc.extract_image(xref)
                        if not base_img:
                            continue
                        img_bytes = base_img['image']
                        img_ext = base_img['ext']
                        w = base_img.get('width', 0)
                        h = base_img.get('height', 0)

                        # Skip tiny images (icons, bullets, spacers)
                        if w < self.MIN_IMAGE_DIM or h < self.MIN_IMAGE_DIM:
                            continue

                        pil_img = Image.open(io.BytesIO(img_bytes))

                        # Resize oversized images to save memory
                        if max(w, h) > self.MAX_IMAGE_PIXELS:
                            ratio = self.MAX_IMAGE_PIXELS / max(w, h)
                            pil_img = pil_img.resize(
                                (int(w * ratio), int(h * ratio)),
                                Image.LANCZOS
                            )
                            buf = io.BytesIO()
                            save_fmt = 'PNG' if img_ext == 'png' else 'JPEG'
                            pil_img.save(buf, format=save_fmt, quality=85)
                            img_bytes = buf.getvalue()
                            w, h = pil_img.size

                        # Get image position on page
                        position = self._get_image_position(page, xref)

                        # Find nearby text (caption detection)
                        caption = self._find_caption(page, position)

                        # Run OCR -- Mistral first, Tesseract fallback
                        ocr_text = ''
                        ocr_details = []
                        if self.mistral_client:
                            ocr_text, ocr_details = self._mistral_ocr_image(
                                pil_img, img_id=f"{doc_id}_p{page.number + 1}_img{img_idx}")
                        if not ocr_text and TESSERACT_AVAILABLE:
                            ocr_text, ocr_details = self._tesseract_ocr_image(pil_img)

                        # Also get any text PyMuPDF detects overlaying the image area
                        overlay_text = ''
                        if position:
                            rect = fitz.Rect(position['x0'], position['y0'],
                                             position['x1'], position['y1'])
                            overlay_text = page.get_text('text', clip=rect).strip()

                        # Save image to disk
                        img_id = f"{doc_id}_p{page.number + 1}_img{img_idx}"
                        img_filename = f"{img_id}.{img_ext}"
                        img_path = os.path.join(
                            app.config['IMAGE_FOLDER'], img_filename)
                        with open(img_path, 'wb') as f:
                            f.write(img_bytes)

                        # Encode as base64 (for inline embedding)
                        b64 = base64.b64encode(img_bytes).decode('utf-8')

                        img_data = {
                            'id': img_id,
                            'page': page.number + 1,
                            'index_on_page': img_idx,
                            'format': img_ext,
                            'width': w,
                            'height': h,
                            'size_bytes': len(img_bytes),
                            'position': position,
                            'caption': caption,
                            'overlay_text': overlay_text,
                            'ocr_text': ocr_text,
                            'ocr_details': ocr_details,
                            'image_file': img_filename,
                            'base64': b64,
                        }
                        self.images.append(img_data)
                        extracted += 1

                    except Exception as e:
                        log.warning(f"Image extraction error on page "
                                    f"{page.number + 1}, img {img_idx}: {e}")

        ocr_engine = 'Mistral AI' if self.mistral_client else ('Tesseract' if TESSERACT_AVAILABLE else 'none')
        log.info(f"Extracted {extracted} images (OCR engine: {ocr_engine})")

    def _get_image_position(self, page, xref):
        """Find the bounding box of an image on a page."""
        try:
            for img_block in page.get_text('dict')['blocks']:
                if img_block.get('type') == 1:  # image block
                    bbox = img_block.get('bbox')
                    if bbox:
                        return {
                            'x0': round(bbox[0], 2),
                            'y0': round(bbox[1], 2),
                            'x1': round(bbox[2], 2),
                            'y1': round(bbox[3], 2),
                        }
        except Exception:
            pass
        return None

    def _find_caption(self, page, position):
        """Try to find a figure caption near the image."""
        if not position:
            return ''
        try:
            # Look for text immediately below the image
            pw = page.rect.width
            below_rect = fitz.Rect(
                0, position['y1'],
                pw, min(position['y1'] + 60, page.rect.height)
            )
            text_below = page.get_text('text', clip=below_rect).strip()

            # Check if it looks like a caption
            caption_patt = re.compile(
                r'^(fig(?:ure)?|table|diagram|chart|image|illustration)'
                r'[\s.:]+',
                re.IGNORECASE
            )
            for line in text_below.split('\n'):
                line = line.strip()
                if caption_patt.match(line):
                    return line
            # Also check text above the image
            above_rect = fitz.Rect(
                0, max(position['y0'] - 40, 0),
                pw, position['y0']
            )
            text_above = page.get_text('text', clip=above_rect).strip()
            for line in text_above.split('\n'):
                line = line.strip()
                if caption_patt.match(line):
                    return line
        except Exception:
            pass
        return ''

    # -----------------------------------------------------------------
    # TESSERACT OCR ENGINE  (fallback when Mistral is not available)
    # -----------------------------------------------------------------

    def _tesseract_ocr_image(self, pil_img):
        """Tesseract fallback: multi-mode OCR on an image for text extraction,
        especially for diagrams with labels, numbers, and axis values."""
        all_text_segments = []
        details = []

        # Convert to RGB if necessary
        if pil_img.mode not in ('RGB', 'L'):
            pil_img = pil_img.convert('RGB')

        configs = [
            ('--psm 3', 'auto'),       # fully automatic
            ('--psm 6', 'block'),      # uniform text block
            ('--psm 11', 'sparse'),    # sparse text -- good for diagrams
        ]

        # Also try preprocessed version
        preprocessed = self._preprocess_for_ocr(pil_img)

        for config_str, mode_name in configs:
            for img_variant, variant_name in [
                (pil_img, 'original'), (preprocessed, 'enhanced')
            ]:
                try:
                    text = pytesseract.image_to_string(
                        img_variant, config=config_str).strip()
                    if text and len(text) > 2:
                        all_text_segments.append(text)
                        details.append({
                            'mode': mode_name,
                            'variant': variant_name,
                            'text': text,
                            'length': len(text),
                        })
                except Exception:
                    pass

        # Also extract structured word-level data
        try:
            data = pytesseract.image_to_data(
                pil_img, output_type=pytesseract.Output.DICT)
            high_conf_words = []
            for i in range(len(data['text'])):
                conf = int(data['conf'][i])
                word = data['text'][i].strip()
                if conf > 50 and word:
                    high_conf_words.append({
                        'text': word,
                        'confidence': conf,
                        'x': data['left'][i],
                        'y': data['top'][i],
                        'w': data['width'][i],
                        'h': data['height'][i],
                    })
            if high_conf_words:
                details.append({
                    'mode': 'word_level',
                    'variant': 'structured',
                    'words': high_conf_words,
                    'word_count': len(high_conf_words),
                })
        except Exception:
            pass

        # Merge: pick the longest unique result
        combined = self._merge_ocr_results(all_text_segments)
        return combined, details

    def _preprocess_for_ocr(self, pil_img):
        """Enhance image contrast & sharpness for better OCR on diagrams."""
        try:
            img = pil_img.convert('L')  # grayscale
            img = ImageEnhance.Contrast(img).enhance(2.0)
            img = ImageEnhance.Sharpness(img).enhance(2.0)
            # Binarize
            img = img.point(lambda x: 0 if x < 140 else 255, '1')
            return img
        except Exception:
            return pil_img

    def _merge_ocr_results(self, segments):
        """Deduplicate and combine OCR results from multiple passes."""
        if not segments:
            return ''
        # Take the longest result (usually most complete)
        segments.sort(key=len, reverse=True)
        best = segments[0]
        # Add any unique lines from other results
        best_lines = set(l.strip().lower() for l in best.split('\n') if l.strip())
        extras = []
        for seg in segments[1:]:
            for line in seg.split('\n'):
                if line.strip() and line.strip().lower() not in best_lines:
                    extras.append(line.strip())
                    best_lines.add(line.strip().lower())
        if extras:
            best += '\n--- additional OCR results ---\n' + '\n'.join(extras)
        return best

    # -----------------------------------------------------------------
    # TEXT FILE EXTRACTION
    # -----------------------------------------------------------------

    def _extract_text_file(self):
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        text = ''
        for enc in encodings:
            try:
                with open(self.filepath, 'r', encoding=enc) as f:
                    text = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        # Split into virtual "pages" for very large files (book-like handling)
        lines = text.split('\n')
        page_size = 80  # lines per virtual page
        vpage = 1
        for i in range(0, len(lines), page_size):
            chunk = '\n'.join(lines[i:i + page_size])
            self.pages.append({
                'page_number': vpage,
                'text': chunk.strip(),
                'char_count': len(chunk.strip()),
                'tables': [],
            })
            vpage += 1
        self.metadata['total_pages'] = len(self.pages)

    # -----------------------------------------------------------------
    # DOCX EXTRACTION
    # -----------------------------------------------------------------

    def _extract_docx(self):
        if not DOCX_AVAILABLE:
            raise ImportError('python-docx not installed')
        doc = DocxDocument(self.filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = '\n\n'.join(paragraphs)

        tables = []
        for table in doc.tables:
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(table_data)

        self.pages = [{
            'page_number': 1,
            'text': full_text.strip(),
            'char_count': len(full_text.strip()),
            'tables': tables,
        }]
        self.metadata['total_pages'] = 1
        self.metadata['paragraph_count'] = len(paragraphs)
        self.metadata['table_count'] = len(tables)

    # -----------------------------------------------------------------
    # TEXT CLEANING
    # -----------------------------------------------------------------

    def _clean_all_pages(self):
        """Deep-clean text on every page."""
        for p in self.pages:
            p['text'] = self._clean_text(p['text'])
            p['char_count'] = len(p['text'])

    def _clean_text(self, text):
        if not text:
            return ''
        # Fix common PDF artefacts
        text = re.sub(r'\x00', '', text)                  # null bytes
        text = re.sub(r'\r\n?', '\n', text)               # normalise newlines
        text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)      # rejoin hyphenated words
        text = re.sub(r'[ \t]+', ' ', text)               # collapse spaces
        text = re.sub(r' \n', '\n', text)                 # trailing space
        text = re.sub(r'\n{4,}', '\n\n\n', text)          # excess blank lines
        text = re.sub(r'(\S)\n(\S)', r'\1 \2', text)      # join broken lines in paragraphs
        # But preserve paragraph separators (double newlines)
        text = re.sub(r'\n\n\n+', '\n\n', text)
        return text.strip()

    # -----------------------------------------------------------------
    # HEADER / FOOTER / PAGE NUMBER REMOVAL
    # -----------------------------------------------------------------

    def _strip_headers_footers(self):
        """Detect repeated first/last lines across pages and remove them."""
        first_lines, last_lines = [], []
        page_num_re = re.compile(r'^\s*[-~]?\s*\d{1,5}\s*[-~]?\s*$')

        for p in self.pages:
            lines = p['text'].split('\n')
            if lines:
                first_lines.append(lines[0].strip())
                last_lines.append(lines[-1].strip())

        threshold = len(self.pages) * 0.3
        first_counts = Counter(first_lines)
        last_counts = Counter(last_lines)

        headers = {t for t, c in first_counts.items() if c > threshold and t}
        footers = {t for t, c in last_counts.items() if c > threshold and t}

        removed = 0
        for p in self.pages:
            lines = p['text'].split('\n')
            cleaned = []
            for idx, line in enumerate(lines):
                s = line.strip()
                # Skip detected headers (first line)
                if idx == 0 and s in headers:
                    removed += 1
                    continue
                # Skip detected footers (last line)
                if idx == len(lines) - 1 and s in footers:
                    removed += 1
                    continue
                # Skip standalone page numbers
                if page_num_re.match(s):
                    removed += 1
                    continue
                cleaned.append(line)
            p['text'] = '\n'.join(cleaned).strip()
            p['char_count'] = len(p['text'])

        if removed:
            log.info(f"Removed {removed} header/footer/page-number lines")

    # -----------------------------------------------------------------
    # STRUCTURE DETECTION (chapters, sections)
    # -----------------------------------------------------------------

    def _detect_structure(self):
        chapters = []
        sections = []

        for p in self.pages:
            for line in p['text'].split('\n')[:5]:  # check first 5 lines
                line = line.strip()
                if not line:
                    continue
                for patt in self.CHAPTER_PATTERNS:
                    if patt.match(line):
                        chapters.append({
                            'title': line,
                            'page': p['page_number'],
                        })
                        break

        # Font-based structure detection (PyMuPDF dict blocks)
        if self.ext == '.pdf':
            try:
                with fitz.open(self.filepath) as doc:
                    font_sizes = []
                    for page in doc:
                        blocks = page.get_text('dict')['blocks']
                        for b in blocks:
                            if b.get('type') == 0:  # text block
                                for span_line in b.get('lines', []):
                                    for span in span_line.get('spans', []):
                                        font_sizes.append(span.get('size', 12))

                    if font_sizes:
                        median_size = sorted(font_sizes)[len(font_sizes) // 2]
                        heading_threshold = median_size * 1.3

                        for page in doc:
                            blocks = page.get_text('dict')['blocks']
                            for b in blocks:
                                if b.get('type') == 0:
                                    for span_line in b.get('lines', []):
                                        for span in span_line.get('spans', []):
                                            sz = span.get('size', 12)
                                            txt = span.get('text', '').strip()
                                            if (sz >= heading_threshold
                                                    and txt
                                                    and len(txt) < 120):
                                                sections.append({
                                                    'text': txt,
                                                    'page': page.number + 1,
                                                    'font_size': round(sz, 1),
                                                })
            except Exception as e:
                log.warning(f"Font-based structure detection failed: {e}")

        self.structure['chapters'] = chapters
        self.structure['sections'] = sections

    # -----------------------------------------------------------------
    # VALIDATION
    # -----------------------------------------------------------------

    def _validate(self):
        """Score overall extraction quality and flag issues."""
        total_chars = sum(p['char_count'] for p in self.pages)
        total_words = sum(len(p['text'].split()) for p in self.pages)
        empty_pages = sum(1 for p in self.pages if p['char_count'] == 0)
        non_empty = len(self.pages) - empty_pages

        page_scores = [self._quality_score(p['text']) for p in self.pages]
        avg_score = sum(page_scores) / len(page_scores) if page_scores else 0

        engines = []
        if self.ext == '.pdf':
            engines = ['pymupdf', 'pdfplumber']
            if self.mistral_client:
                engines.append('mistral-ocr')
        else:
            engines = ['native']

        self.quality_report.update({
            'overall_score': round(avg_score, 4),
            'total_characters': total_chars,
            'total_words': total_words,
            'total_pages': len(self.pages),
            'empty_pages': empty_pages,
            'non_empty_pages': non_empty,
            'total_images': len(self.images),
            'ocr_engine': 'mistral' if self.mistral_client else ('tesseract' if TESSERACT_AVAILABLE else 'none'),
            'ocr_available': OCR_AVAILABLE,
            'images_with_ocr': sum(1 for img in self.images if img.get('ocr_text')),
            'images_with_captions': sum(1 for img in self.images if img.get('caption')),
            'chapters_detected': len(self.structure.get('chapters', [])),
            'sections_detected': len(self.structure.get('sections', [])),
            'extraction_passes': (4 if self.mistral_client else 3) if self.ext == '.pdf' else 1,
            'engines_used': engines,
        })

        # Warnings
        warnings = []
        if empty_pages > len(self.pages) * 0.2:
            warnings.append(f'{empty_pages} pages had no extractable text '
                            '(possibly scanned/image-only pages)')
        if avg_score < 0.3:
            warnings.append('Overall extraction quality is low. '
                            'The document may be scanned or encrypted.')
        if not self.mistral_client and not TESSERACT_AVAILABLE and self.images:
            warnings.append('No OCR engine available. Provide a Mistral API '
                            'key or install Tesseract for image/diagram OCR.')
        self.quality_report['warnings'] = warnings

    def _ai_validate(self):
        """Uses OpenAI and/or Sarvam AI concurrently to validate full extraction quality."""
        if not OPENAI_AVAILABLE:
            return

        openai_key = self.options.get('openai_api_key')
        sarvam_key = self.options.get('sarvam_api_key')

        if not openai_key and not sarvam_key:
            return

        # Prepare the full text
        full_text = '\n'.join(p['text'] for p in self.pages if p['text'])
        if not full_text:
            return

        # Cap at a very large safe limit (approx 250,000 tokens) to prevent outright API rejection,
        # but large enough to cover the vast majority of entire books/documents fully.
        sample = full_text[:1000000]

        prompt = (
            "You are an expert AI extraction reviewer. Evaluate the following complete extracted document "
            "for overall OCR quality, layout preservation, missing spaces, broken words, or garbled characters. "
            "Reply strictly in JSON format with exactly 2 keys: 'score' (a number from 1 to 10) and "
            "'reason' (a short 1-sentence reason for the score). Do not include any other text.\n\n"
            f"--- TEXT ---\n{sample}\n--- END TEXT ---"
        )
        responses = {}

        def _run_openai():
            try:
                client = OpenAI(api_key=openai_key)
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    response_format={"type": "json_object"},
                    temperature=0.0
                )
                res_content = resp.choices[0].message.content
                data = json.loads(res_content)
                log.info(f"OpenAI validation complete: score={data.get('score')}")
                return ('openai', data)
            except Exception as e:
                log.warning(f"OpenAI validation failed: {e}")
                return ('openai', {'error': str(e)})

        def _run_sarvam():
            try:
                client = OpenAI(base_url="https://api.sarvam.ai/v1/", api_key=sarvam_key)
                resp = client.chat.completions.create(
                    model="sarvam-30b",
                    messages=[
                        {"role": "system", "content": "You are a helpful AI assistant that outputs only JSON."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.0
                )
                res_content = resp.choices[0].message.content.strip()
                if res_content.startswith('```json'):
                    res_content = res_content[7:-3].strip()
                elif res_content.startswith('```'):
                    res_content = res_content[3:-3].strip()
                data = json.loads(res_content)
                log.info(f"Sarvam AI validation complete: score={data.get('score')}")
                return ('sarvam', data)
            except Exception as e:
                log.warning(f"Sarvam AI validation failed: {e}")
                return ('sarvam', {'error': str(e)})

        # Run simultaneously
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            futures = []
            if openai_key:
                futures.append(executor.submit(_run_openai))
            if sarvam_key:
                futures.append(executor.submit(_run_sarvam))

            for future in concurrent.futures.as_completed(futures):
                key, result = future.result()
                responses[key] = result

        if responses:
            self.quality_report['ai_validation'] = responses

    # -----------------------------------------------------------------
    # QUALITY SCORING
    # -----------------------------------------------------------------

    def _quality_score(self, text):
        """Score text quality 0-1.  Higher = better extraction."""
        if not text:
            return 0.0
        words = text.split()
        if not words:
            return 0.0

        # Factor 1: word count (more is usually better).
        # We increase the saturation point to 600 words to reward more complete text.
        wc_score = min(len(words) / 600, 1.0)

        # Factor 2: alphabetic character ratio
        alpha = sum(1 for c in text if c.isalpha())
        total = len(text)
        alpha_r = alpha / total if total else 0

        # Factor 3: mean word length (garbled text has very short words)
        avg_wl = sum(len(w) for w in words) / len(words)
        wl_score = min(avg_wl / 5, 1.0) if avg_wl < 15 else 0.5

        # Factor 4: newline density (too high = formatting noise)
        nl_ratio = text.count('\n') / total if total else 0
        nl_score = max(1.0 - nl_ratio * 10, 0)

        # Weighted average
        return (wc_score * 0.35 + alpha_r * 0.30
                + wl_score * 0.15 + nl_score * 0.20)


# =====================================================================
# TEXT UTILITIES
# =====================================================================

def split_into_sentences(text):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sentences if s.strip()]


def split_into_paragraphs(text):
    paragraphs = re.split(r'\n\s*\n', text)
    return [p.strip() for p in paragraphs if p.strip()]


def chunk_text(text, chunk_size=512, overlap=64):
    """Split text into overlapping chunks (word-based)."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(' '.join(words[start:end]))
        if end >= len(words):
            break
        start += chunk_size - overlap
    return chunks


def clean_for_output(text):
    """Final text cleanup before writing to JSON."""
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# =====================================================================
# OUTPUT FORMAT HANDLERS
# =====================================================================

def format_raw(result, filename, options):
    full_text = '\n\n'.join(p['text'] for p in result['pages'] if p['text'])
    full_text = clean_for_output(full_text)

    # Strip ai_validation from quality report
    quality = {k: v for k, v in result['quality'].items() if k != 'ai_validation'}

    return {
        'id': str(uuid.uuid4()),
        'source_file': filename,
        'format': 'raw_document',
        'created_at': datetime.now(timezone.utc).isoformat(),
        'metadata': result['metadata'],
        'quality_report': quality,
        'structure': result['structure'],
        'content': {
            'full_text': full_text,
            'total_characters': len(full_text),
            'total_words': len(full_text.split()),
            'pages': [{
                'page_number': p['page_number'],
                'text': p['text'],
                'char_count': p['char_count'],
                'tables': p.get('tables', []),
            } for p in result['pages']],
        },
        'checksum': hashlib.md5(full_text.encode()).hexdigest(),
    }


def format_chunks(result, filename, options):
    full_text = '\n\n'.join(p['text'] for p in result['pages'] if p['text'])
    full_text = clean_for_output(full_text)
    cs = options.get('chunk_size', 512)
    co = options.get('chunk_overlap', 64)
    chunks = chunk_text(full_text, chunk_size=cs, overlap=co)

    quality = {k: v for k, v in result['quality'].items() if k != 'ai_validation'}

    return {
        'source_file': filename,
        'format': 'chunked',
        'created_at': datetime.now(timezone.utc).isoformat(),
        'config': {'chunk_size': cs, 'chunk_overlap': co, 'unit': 'words'},
        'quality_report': quality,
        'total_chunks': len(chunks),
        'chunks': [{
            'id': f'{filename}::chunk_{i}',
            'index': i,
            'text': c,
            'word_count': len(c.split()),
            'char_count': len(c),
        } for i, c in enumerate(chunks)],
    }


def format_paragraphs(result, filename, options):
    full_text = '\n\n'.join(p['text'] for p in result['pages'] if p['text'])
    full_text = clean_for_output(full_text)
    paras = split_into_paragraphs(full_text)

    quality = {k: v for k, v in result['quality'].items() if k != 'ai_validation'}

    return {
        'source_file': filename,
        'format': 'paragraphs',
        'created_at': datetime.now(timezone.utc).isoformat(),
        'quality_report': quality,
        'total_paragraphs': len(paras),
        'data': [{
            'id': f'{filename}::para_{i}',
            'index': i,
            'text': p,
            'word_count': len(p.split()),
        } for i, p in enumerate(paras)],
    }


def format_sentences(result, filename, options):
    full_text = '\n\n'.join(p['text'] for p in result['pages'] if p['text'])
    full_text = clean_for_output(full_text)
    sents = split_into_sentences(full_text)

    quality = {k: v for k, v in result['quality'].items() if k != 'ai_validation'}

    return {
        'source_file': filename,
        'format': 'sentences',
        'created_at': datetime.now(timezone.utc).isoformat(),
        'quality_report': quality,
        'total_sentences': len(sents),
        'data': [{
            'id': f'{filename}::sent_{i}',
            'index': i,
            'text': s,
            'word_count': len(s.split()),
        } for i, s in enumerate(sents)],
    }


def format_jsonl(result, filename, options):
    full_text = '\n\n'.join(p['text'] for p in result['pages'] if p['text'])
    full_text = clean_for_output(full_text)
    cs = options.get('chunk_size', 512)
    chunks = chunk_text(full_text, chunk_size=cs, overlap=0)

    lines = []
    for i, c in enumerate(chunks):
        lines.append(json.dumps({
            'text': c,
            'source': filename,
            'index': i,
        }, ensure_ascii=False))

    return '\n'.join(lines)


def format_conversation(result, filename, options):
    full_text = '\n\n'.join(p['text'] for p in result['pages'] if p['text'])
    full_text = clean_for_output(full_text)
    cs = options.get('chunk_size', 1024)
    chunks = chunk_text(full_text, chunk_size=cs, overlap=0)

    system_prompt = options.get(
        'system_prompt',
        'You are a knowledgeable assistant. Answer based on the provided content.')

    quality = {k: v for k, v in result['quality'].items() if k != 'ai_validation'}

    conversations = []
    for i, c in enumerate(chunks):
        conversations.append({
            'id': f'{filename}::conv_{i}',
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user',
                 'content': f'What information is in section {i + 1} of {filename}?'},
                {'role': 'assistant', 'content': c},
            ],
        })

    return {
        'source_file': filename,
        'format': 'conversation',
        'created_at': datetime.now(timezone.utc).isoformat(),
        'quality_report': quality,
        'total_conversations': len(conversations),
        'data': conversations,
    }


# =====================================================================
# TEXTBOOK FORMAT HANDLER -- Structured Textbook JSON via Mistral AI
# =====================================================================

TEXTBOOK_SCHEMA_PROMPT = """You are a textbook content extractor. Parse ALL content into the JSON schema below. Return ONLY valid JSON, no markdown fences.

RULES:
1. "content" is the #1 priority field: capture FULL teaching text for each section. NEVER truncate.
2. "topics": only named concepts (e.g. "Photosynthesis", "Ohm's Law"). NOT random words.
3. EXERCISES: extract ALL questions by type. NEVER include answers in question arrays. Answers go in answer_keys.
4. Use exact section numbering from textbook. Sub-sections (1.4.1) go under parent (1.4).
5. "count" must match array length for all exercise types.

JSON SCHEMA:
{
  "chapters": [{
    "chapter_number": 1,
    "chapter_name": "<full title as printed>",
    "sections": [{
      "section_number": "1.1",
      "section_name": "<heading>",
      "content": "<FULL teaching text - ALL paragraphs, definitions, explanations>",
      "sub_sections": [{"sub_section_number":"1.1.1","sub_section_name":"","content":"<full text>","topics":[],"formulas":[],"key_terms":[{"term":"","definition":""}],"diagrams":[{"figure_id":"","description":"","labels":[]}],"chemical_equations":[],"derivations":[{"title":"","steps":""}],"worked_examples":[{"problem":"","solution":""}]}],
      "topics": [],
      "formulas": ["<formula with units>"],
      "key_terms": [{"term":"","definition":""}],
      "diagrams": [{"figure_id":"Fig 1.1","description":"what it shows","labels":["part1"]}],
      "chemical_equations": [],
      "derivations": [{"title":"","steps":"<full derivation>"}],
      "worked_examples": [{"problem":"","solution":""}]
    }],
    "textbook_page_start": null, "textbook_page_end": null,
    "introduction": "<intro paragraphs>",
    "exercises": {
      "mcq":{"count":0,"questions":[{"question":"","options":"(A)...(B)...(C)...(D)..."}]},
      "assertion_reason":{"count":0,"questions":[]},
      "fill_in_the_blanks":{"count":0,"questions":[]},
      "true_false":{"count":0,"questions":[]},
      "match_the_following":{"count":0,"questions":[]},
      "very_short_answer":{"count":0,"questions":[]},
      "short_answer":{"count":0,"questions":[]},
      "long_answer":{"count":0,"questions":[]},
      "numerical_problems":{"count":0,"questions":[]},
      "diagram_based":{"count":0,"questions":[]},
      "hots":{"count":0,"questions":[]},
      "activity_based":{"count":0,"questions":[]}
    },
    "answer_keys":{"exercise_answers":[]},
    "learning_objectives": [],
    "chapter_summary":{"key_definitions":[],"important_laws":[],"important_reactions":[],"nature_points":[],"importance_points":[]}
  }]
}

Extract EVERY section, ALL exercises, ALL formulas/equations/derivations. NEVER skip content."""



def _call_mistral_for_structuring(mistral_client, text_chunk, system_prompt=None, prompt_prefix=''):
    """Send text to Mistral AI for structured textbook parsing.
    Includes retry logic for robustness."""
    if system_prompt is None:
        system_prompt = TEXTBOOK_SCHEMA_PROMPT

    user_content = f'{prompt_prefix}\n\nHere is the extracted textbook text:\n\n{text_chunk}' if prompt_prefix else f'Here is the extracted textbook text:\n\n{text_chunk}'

    max_retries = 2
    for attempt in range(max_retries + 1):
        try:
            resp = mistral_client.chat.complete(
                model='mistral-large-latest',
                messages=[
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': user_content}
                ],
                response_format={'type': 'json_object'},
                temperature=0.0,
                max_tokens=16384,
            )
            content = resp.choices[0].message.content.strip()
            # Try to fix truncated JSON
            if content and not content.endswith('}'):
                # Try adding missing closing braces
                open_braces = content.count('{') - content.count('}')
                open_brackets = content.count('[') - content.count(']')
                content += ']' * max(open_brackets, 0) + '}' * max(open_braces, 0)
            parsed = json.loads(content)
            log.info(f'Mistral structuring OK: {len(content)} chars, {len(str(parsed))} parsed')
            return parsed
        except json.JSONDecodeError as e:
            log.warning(f'Mistral returned invalid JSON (attempt {attempt + 1}): {e}')
            if attempt < max_retries:
                continue
            return None
        except Exception as e:
            if attempt < max_retries:
                log.warning(f'Mistral structuring attempt {attempt + 1} failed: {e}, retrying...')
                continue
            log.warning(f'Mistral structuring failed after {max_retries + 1} attempts: {e}')
            return None


def _detect_subject(pages):
    """Auto-detect the subject of the textbook/chapter by scanning content for keywords."""
    # Sample first 5 pages + last 2 pages for keyword density
    sample_pages = pages[:5] + pages[-2:] if len(pages) > 5 else pages
    sample_text = ' '.join(p.get('text', '') for p in sample_pages).lower()

    subject_signatures = {
        'Biology': {
            'keywords': [
                'cell', 'organism', 'species', 'plant', 'animal', 'tissue', 'organ',
                'photosynthesis', 'respiration', 'dna', 'rna', 'gene', 'chromosome',
                'mitosis', 'meiosis', 'evolution', 'ecosystem', 'habitat', 'taxonomy',
                'kingdom', 'phylum', 'class', 'genus', 'biodiversity', 'flora', 'fauna',
                'enzyme', 'protein', 'nucleus', 'cytoplasm', 'membrane', 'chloroplast',
                'living organism', 'biological', 'anatomy', 'physiology', 'morphology',
                'reproduction', 'heredity', 'ecology', 'digestion', 'excretion',
            ],
            'strong': ['photosynthesis', 'mitosis', 'meiosis', 'chromosome', 'dna', 'kingdom', 'phylum', 'taxonomy'],
        },
        'Physics': {
            'keywords': [
                'force', 'velocity', 'acceleration', 'mass', 'energy', 'momentum',
                'newton', 'gravity', 'friction', 'wave', 'frequency', 'wavelength',
                'electric', 'magnetic', 'current', 'voltage', 'resistance', 'circuit',
                'optics', 'lens', 'mirror', 'refraction', 'reflection', 'spectrum',
                'thermodynamics', 'pressure', 'temperature', 'heat', 'work', 'power',
                'kinetic energy', 'potential energy', 'joule', 'watt', 'ampere',
                'magnetic field', 'electromagnetic', 'nuclear', 'radioactive',
                'displacement', 'torque', 'angular', 'rotational', 'oscillation',
            ],
            'strong': ['newton', 'electromagnetic', 'thermodynamics', 'refraction', 'kinetic energy', 'ohm'],
        },
        'Chemistry': {
            'keywords': [
                'atom', 'molecule', 'element', 'compound', 'reaction', 'chemical',
                'acid', 'base', 'salt', 'solution', 'solvent', 'solute', 'mole',
                'periodic table', 'electron', 'proton', 'neutron', 'orbital',
                'bond', 'covalent', 'ionic', 'metallic', 'oxidation', 'reduction',
                'catalyst', 'equilibrium', 'concentration', 'titration', 'ph',
                'organic', 'inorganic', 'polymer', 'hydrocarbon', 'alkane', 'alkene',
                'iupac', 'nomenclature', 'valence', 'electronegativity', 'redox',
                'stoichiometry', 'molarity', 'reagent', 'precipitate',
            ],
            'strong': ['periodic table', 'covalent', 'ionic', 'stoichiometry', 'iupac', 'electronegativity'],
        },
        'Mathematics': {
            'keywords': [
                'equation', 'theorem', 'proof', 'formula', 'calculate', 'solve',
                'triangle', 'circle', 'angle', 'parallel', 'perpendicular',
                'polynomial', 'quadratic', 'linear', 'matrix', 'determinant',
                'integral', 'derivative', 'differentiation', 'integration', 'limit',
                'trigonometry', 'sine', 'cosine', 'tangent', 'logarithm',
                'probability', 'statistics', 'mean', 'median', 'variance',
                'geometry', 'algebra', 'arithmetic', 'coordinate', 'vector',
                'function', 'domain', 'range', 'continuous', 'sequence', 'series',
            ],
            'strong': ['theorem', 'polynomial', 'quadratic', 'trigonometry', 'integral', 'derivative'],
        },
        'Statistics': {
            'keywords': [
                'index number', 'correlation', 'regression', 'frequency', 'distribution',
                'mean', 'median', 'mode', 'standard deviation', 'variance',
                'probability', 'sample', 'population', 'hypothesis', 'scatter diagram',
                'time series', 'trend', 'seasonal', 'cyclical', 'moving average',
            ],
            'strong': ['correlation', 'regression', 'index number', 'scatter diagram', 'time series'],
        },
    }

    scores = {}
    for subject, data in subject_signatures.items():
        score = 0
        for kw in data['keywords']:
            count = sample_text.count(kw)
            if count > 0:
                score += min(count, 5)  # Cap at 5 per keyword
        # Strong indicators get bonus
        for kw in data.get('strong', []):
            if kw in sample_text:
                score += 10
        scores[subject] = score

    if not scores:
        return ''

    best_subject = max(scores, key=scores.get)
    best_score = scores[best_subject]

    # Need at least some minimum confidence
    if best_score < 5:
        return ''

    log.info(f'Auto-detected subject: {best_subject} (score={best_score}, '
             f'runners-up: {sorted(((s, sc) for s, sc in scores.items() if s != best_subject), key=lambda x: -x[1])[:2]})')
    return best_subject


def _detect_document_type(full_text, pages):
    """Detect whether the document is a full book or a single chapter.
    Returns ('single_chapter', chapter_info) or ('full_book', None)."""
    page_count = len(pages)

    # Count chapter/unit boundaries
    chapter_boundary_re = re.compile(
        r'^(?:---\s*Page\s+\d+\s*---\s*\n\s*)?(?:chapter|ch\.?|unit)\s+(\d+|[ivxlc]+)[\s.:]+(.+)',
        re.IGNORECASE | re.MULTILINE
    )
    chapter_matches = list(chapter_boundary_re.finditer(full_text))

    # Check for table of contents
    toc_indicators = ['table of contents', 'contents', 'index']
    first_pages_text = ' '.join(p.get('text', '') for p in pages[:3]).lower()
    has_toc = any(ind in first_pages_text for ind in toc_indicators)

    # Detect chapter name from the first chapter marker if only 1
    chapter_info = None
    if chapter_matches:
        chapter_info = {
            'number': chapter_matches[0].group(1).strip(),
            'name': chapter_matches[0].group(2).strip(),
        }

    # Decision logic
    if page_count <= 5:
        doc_type = 'single_chapter'
    elif page_count <= 40 and len(chapter_matches) <= 1 and not has_toc:
        doc_type = 'single_chapter'
    elif len(chapter_matches) >= 3:
        doc_type = 'full_book'
    elif has_toc and len(chapter_matches) >= 2:
        doc_type = 'full_book'
    elif page_count > 80:
        doc_type = 'full_book'
    else:
        # Default: treat as single chapter for smaller docs
        doc_type = 'single_chapter' if page_count <= 50 else 'full_book'

    log.info(f'Document type: {doc_type} ({page_count} pages, '
             f'{len(chapter_matches)} chapter markers, TOC={has_toc})')
    return doc_type, chapter_info


def _split_text_by_chapters(full_text):
    """Split full document text into per-chapter segments using chapter boundary detection.
    Returns a list of (hint, segment_text, start_page) tuples."""
    # List of regex patterns for chapter/unit/major-section boundaries
    patterns = [
        # Chapter 1: Name or Ch 1. Name  (strongest signal)
        re.compile(r'^(?:---\s*Page\s+\d+\s*---\s*\n\s*)?(?:chapter|ch\.?)\s+(\d+|[ivxlc]+)[\s.:]*(.*)', re.IGNORECASE | re.MULTILINE),
        # Unit 1: Name
        re.compile(r'^(?:---\s*Page\s+\d+\s*---\s*\n\s*)?(?:unit)\s+(\d+|[ivxlc]+)[\s.:]*(.*)', re.IGNORECASE | re.MULTILINE),
        # 1. Title Case Name (requires dot after number and title-case name with 2+ words)
        re.compile(r'^(?:---\s*Page\s+\d+\s*---\s*\n\s*)?(\d+)\.[\s\t]+([A-Z][a-z]+(?:\s+[A-Za-z]+){1,})\s*$', re.MULTILINE),
        # Lesson 1: Name (common in some textbooks)
        re.compile(r'^(?:---\s*Page\s+\d+\s*---\s*\n\s*)?(?:lesson)\s+(\d+|[ivxlc]+)[\s.:]*(.*)', re.IGNORECASE | re.MULTILINE),
    ]

    boundaries = []
    # Find all matches for all patterns
    for patt in patterns:
        for m in patt.finditer(full_text):
            # Extract page number if present in the matched prefix
            match_text = m.group(0)
            page_m = re.search(r'---\s*Page\s+(\d+)\s*---', match_text)
            page_num = int(page_m.group(1)) if page_m else None
            
            chapter_num_str = m.group(1).strip()
            title = m.group(2).strip() if m.group(2) else ''
            
            # Skip if title is empty or is just the subject name repeated (running header)
            if not title or len(title) < 3:
                continue
            
            # Skip if chapter number matches the page number (likely a running header like "4 BIOLOGY")
            if page_num and chapter_num_str.isdigit() and int(chapter_num_str) == page_num:
                continue
                
            hint = f"{chapter_num_str} {title}".strip()
            boundaries.append({'pos': m.start(), 'hint': hint, 'page': page_num})

    # Sort boundaries by position and remove duplicates (overlaps)
    boundaries.sort(key=lambda x: x['pos'])
    
    # Filter out running headers: if the same title appears 3+ times, it's a running header
    title_counts = {}
    for b in boundaries:
        # Normalize title for comparison
        title_part = re.sub(r'^\d+\s*', '', b['hint']).strip().lower()
        title_counts[title_part] = title_counts.get(title_part, 0) + 1
    
    filtered_boundaries = []
    for b in boundaries:
        title_part = re.sub(r'^\d+\s*', '', b['hint']).strip().lower()
        if title_counts.get(title_part, 0) >= 3:
            # This title appears too many times - it's a running header, skip all of them
            continue
        filtered_boundaries.append(b)
    
    unique_boundaries = []
    last_pos = -1
    for b in filtered_boundaries:
        if b['pos'] > last_pos + 10: # avoid very close matches
            unique_boundaries.append(b)
            last_pos = b['pos']

    if not unique_boundaries:
        return [('Full Document', full_text, 1)]

    segments = []
    for i, b in enumerate(unique_boundaries):
        start = b['pos']
        end = unique_boundaries[i + 1]['pos'] if i + 1 < len(unique_boundaries) else len(full_text)
        text = full_text[start:end].strip()
        if text:
            segments.append((b['hint'], text, b['page']))

    # Handle preamble
    if unique_boundaries[0]['pos'] > 300:
        preamble = full_text[:unique_boundaries[0]['pos']].strip()
        if preamble:
            segments.insert(0, ('Front Matter', preamble, 1))

    return segments


def _merge_chapters(master_list, new_chapters):
    """Smart merge of new chapter objects into the master list.
    Handles duplicates by merging fields instead of replacing."""
    if not new_chapters:
        return master_list

    for new_ch in new_chapters:
        # Try to find existing chapter by number or name
        num = new_ch.get('chapter_number')
        name = str(new_ch.get('chapter_name', '')).lower().strip()
        
        existing = None
        for ch in master_list:
            if num is not None and ch.get('chapter_number') == num:
                existing = ch
                break
            if name and name == str(ch.get('chapter_name', '')).lower().strip():
                existing = ch
                break
        
        if not existing:
            master_list.append(new_ch)
            continue
            
        # Merge fields for existing chapter
        # 1. Introduction & Content: concatenate if unique
        new_intro = new_ch.get('introduction', '').strip()
        if new_intro and new_intro not in existing.get('introduction', ''):
            existing['introduction'] = (existing.get('introduction', '') + '\n\n' + new_intro).strip()
            
        # 2. Sections: deduplicate by number/name
        existing_sec_names = {str(s.get('section_name', '')).lower().strip() for s in existing.get('sections', [])}
        for sec in new_ch.get('sections', []):
            s_name = str(sec.get('section_name', '')).lower().strip()
            if s_name not in existing_sec_names:
                existing.setdefault('sections', []).append(sec)
                existing_sec_names.add(s_name)
            else:
                # Merge sub-sections if section already exists
                for target_sec in existing['sections']:
                    if str(target_sec.get('section_name', '')).lower().strip() == s_name:
                        existing_sub_names = {str(ss.get('sub_section_name', '')).lower().strip() for ss in target_sec.get('sub_sections', [])}
                        for sub in sec.get('sub_sections', []):
                            ss_name = str(sub.get('sub_section_name', '')).lower().strip()
                            if ss_name not in existing_sub_names:
                                target_sec.setdefault('sub_sections', []).append(sub)
                                existing_sub_names.add(ss_name)
                        break

        # 3. Exercises: additive merge
        if 'exercises' in new_ch:
            for ex_type, ex_data in new_ch['exercises'].items():
                if not isinstance(ex_data, dict) or 'questions' not in ex_data:
                    continue
                
                target_ex = existing.setdefault('exercises', {}).setdefault(ex_type, {'count': 0, 'questions': []})
                existing_qs = set()
                for q in target_ex.get('questions', []):
                    if isinstance(q, str): existing_qs.add(q.strip().lower())
                    elif isinstance(q, dict): existing_qs.add(str(q.get('question', '')).strip().lower())
                
                for q in ex_data.get('questions', []):
                    q_text = q if isinstance(q, str) else str(q.get('question', ''))
                    if q_text.strip().lower() not in existing_qs:
                        target_ex['questions'].append(q)
                        existing_qs.add(q_text.strip().lower())
                
                target_ex['count'] = len(target_ex['questions'])

        # 4. Learning Objectives: unique extend
        existing_los = {lo.lower().strip() for lo in existing.get('learning_objectives', []) if isinstance(lo, str)}
        for lo in new_ch.get('learning_objectives', []):
            if isinstance(lo, str) and lo.lower().strip() not in existing_los:
                existing.setdefault('learning_objectives', []).append(lo)
                existing_los.add(lo.lower().strip())

        # 5. Page ranges
        if new_ch.get('textbook_page_start'):
            if not existing.get('textbook_page_start') or new_ch['textbook_page_start'] < existing['textbook_page_start']:
                existing['textbook_page_start'] = new_ch['textbook_page_start']
        if new_ch.get('textbook_page_end'):
            if not existing.get('textbook_page_end') or new_ch['textbook_page_end'] > existing['textbook_page_end']:
                existing['textbook_page_end'] = new_ch['textbook_page_end']

    return master_list


def _detect_exercise_pages(pages):
    """Scan pages to identify ranges that likely contain exercises/assessments."""
    exercise_ranges = []
    exercise_keywords = {
        'exercise', 'exercises', 'practice', 'questions', 'problems', 'objective',
        'multiple choice', 'mcq', 'fill in the blank', 'true or false', 'match the',
        'one sentence', 'short questions', 'detailed questions', 'essay type'
    }
    
    current_range = None
    
    for p in pages:
        text = p.get('text', '').lower()
        # Does this page look like an exercise page?
        score = sum(1 for kw in exercise_keywords if kw in text)
        
        # If page has "Exercise 1.1" etc. it's a very strong indicator
        if re.search(r'exercise\s+\d+\.\d+', text) or re.search(r'section\s+exercise', text):
            score += 5
            
        is_exercise = score >= 3
        
        if is_exercise:
            if not current_range:
                current_range = {'start': p['page_number'], 'end': p['page_number'], 'strength': score}
            else:
                current_range['end'] = p['page_number']
                current_range['strength'] = max(current_range['strength'], score)
        else:
            if current_range:
                exercise_ranges.append(current_range)
                current_range = None
                
    if current_range:
        exercise_ranges.append(current_range)
        
    return exercise_ranges


def _clean_exercises(exercises):
    """Post-process exercises to remove answer keys mixed into questions
    and fix counts."""
    # Patterns that indicate an answer rather than a question
    answer_patterns = [
        re.compile(r'^\s*[\(\[]?\d+[\)\]]\s*[\(\[]\s*[a-dA-D]\s*[\)\]]', re.IGNORECASE),  # (1) (a)
        re.compile(r'^\s*r\s*=\s*[-\d.]+'),  # r = 0.81
        re.compile(r'^\s*I\s*[=_]\s*[-\d.]+'),  # I = 123.80
        re.compile(r'^\s*[Ii]ndex\s+numbers?\s*[:=]'),  # Index number = ...
        re.compile(r'^\s*(?:Fixed|Chain)\s+base\s+index\s+numbers?\s*[:=]'),
        re.compile(r'^\s*(?:Real\s+wages?|Purchasing\s+power)\s*[:=]'),
        re.compile(r'^\s*(?:b|by|bx|byx|bxy)\s*=\s*[-\d.]+'),  # regression answers
        re.compile(r'^\s*[ŷ$§]\s*=\s*[-\d.]+'),  # regression line answers
        re.compile(r'^\s*(?:Error|error)\s*[=:]\s*'),
        re.compile(r'^\s*(?:Scale|X-axis|Y-axis|Xeaxis|Yeaxis)\s*:'),  # graph instructions
        re.compile(r'^\s*(?:Three|Four|Five)\s+(?:yearly|monthly|quarterly)\s*[\|]'),  # table answers
        re.compile(r'^\s*[\(\[]?\s*[a-d]\s*[\)\]]?\s*$'),  # standalone (a), (b), etc.
        re.compile(r'^\s*\d+\.\s*\d+\s*$'),  # standalone numbers like "16187"
        re.compile(r'^\s*(?:Corrected|corrected)\s+'),
        re.compile(r'^\s*(?:Both\s+variables|Regression\s+coefficient\s+will)'),  # explanation answers
        # Science-specific answer patterns
        re.compile(r'^\s*(?:Ans|Answer|Sol|Solution)[.:\s]', re.IGNORECASE),  # Ans: ...
        re.compile(r'^\s*=\s*[-\d.]+\s*(?:m|cm|kg|g|J|N|W|V|A|Ω|Hz|mol|L|s|K|°C)'),  # = 42 J
        re.compile(r'^\s*[-\d.]+\s*(?:m|cm|kg|g|J|N|W|V|A|Ω|Hz|mol|L|s|K|°C)\s*$'),  # standalone 42 J
        re.compile(r'^\s*x\s*=\s*[-\d.]+\s*$'),  # x = 5
        re.compile(r'^\s*y\s*=\s*[-\d.]+\s*$'),  # y = 3
        re.compile(r'^\s*(?:Hint|Note)\s*:', re.IGNORECASE),  # Hint: ...
    ]

    def is_answer(text):
        """Check if a text looks like an answer rather than a question."""
        if not text or len(text.strip()) < 3:
            return True
        for patt in answer_patterns:
            if patt.match(text.strip()):
                return True
        return False

    if not isinstance(exercises, dict):
        return exercises

    for ex_type, ex_data in exercises.items():
        if ex_type == 'answer_keys':
            continue
        if not isinstance(ex_data, dict):
            continue

        questions = ex_data.get('questions', [])
        if not isinstance(questions, list):
            continue

        # Clean MCQ questions (list of dicts)
        if ex_type == 'mcq':
            cleaned = []
            for q in questions:
                if isinstance(q, dict):
                    q_text = q.get('question', '')
                    if q_text and not is_answer(q_text) and len(q_text.strip()) > 10:
                        cleaned.append(q)
                elif isinstance(q, str) and not is_answer(q) and len(q.strip()) > 10:
                    cleaned.append({'question': q, 'options': ''})
            ex_data['questions'] = cleaned
            ex_data['count'] = len(cleaned)
        else:
            # Clean string question arrays
            cleaned = []
            for q in questions:
                if isinstance(q, str) and not is_answer(q) and len(q.strip()) > 10:
                    cleaned.append(q.strip())
                elif isinstance(q, dict):
                    q_text = q.get('question', '') or q.get('text', '')
                    if q_text and not is_answer(q_text) and len(q_text.strip()) > 10:
                        cleaned.append(q_text.strip())
            ex_data['questions'] = cleaned
            ex_data['count'] = len(cleaned)

    return exercises


def _clean_topics(topics):
    """Remove garbage topic entries - single words, sentence fragments, etc."""
    if not isinstance(topics, list):
        return []

    # Words that are never valid standalone topics (but NOT scientific terms)
    garbage_words = {
        'there', 'if', 'it', 'the', 'we', 'a', 'an', 'is', 'are', 'was', 'were',
        'this', 'that', 'these', 'those', 'when', 'where', 'which', 'who', 'what',
        'how', 'why', 'but', 'and', 'or', 'not', 'no', 'yes', 'so', 'also', 'hence',
        'thus', 'since', 'because', 'therefore', 'however', 'moreover', 'furthermore',
        'less', 'more', 'here', 'atmost', 'once',
        'another', 'some', 'sometimes', 'unless', 'ata', 'its', 'two', 'three',
        'proper', 'important', 'necessary', 'effect',
    }
    # Scientific terms that MUST be preserved even if short
    science_terms = {
        'dna', 'rna', 'atp', 'adp', 'nad', 'fad', 'ph', 'emf', 'si', 'cgs',
        'iupac', 'redox', 'hcl', 'nacl', 'h2o', 'co2', 'o2', 'n2', 'h2',
        'ohm', 'volt', 'watt', 'joule', 'kelvin', 'mole', 'gene', 'cell',
    }

    cleaned = []
    for topic in topics:
        if not isinstance(topic, str):
            continue
        t = topic.strip()

        # Skip empty or very short entries (but allow short science terms)
        if len(t) < 2:
            continue
        if len(t) < 4 and t.lower() not in science_terms:
            continue

        # Skip single-word garbage (but NOT science terms)
        if t.lower().rstrip('.,;:') in garbage_words and t.lower().rstrip('.,;:') not in science_terms:
            continue

        # Skip entries that are just 1-2 common words
        words = t.split()
        if len(words) <= 2 and all(w.lower().rstrip('.,;:') in garbage_words for w in words):
            continue

        # Skip entries that look like sentence fragments (start with lowercase, end abruptly)
        # But preserve known science terms like "pH"
        if (t[0].islower() and not any(c in t for c in '.?!') and len(t) < 30
                and not re.match(r'^[a-z]+\s+[A-Z]', t)
                and t.lower() not in science_terms):
            continue

        # Skip entries that are clearly OCR garbage (mostly symbols/numbers)
        # But allow chemical formulas (H₂SO₄, NaCl) which have low alpha ratio
        alpha_ratio = sum(1 for c in t if c.isalpha()) / max(len(t), 1)
        has_subscript = any(c in t for c in '₀₁₂₃₄₅₆₇₈₉')
        if alpha_ratio < 0.2 and not has_subscript:
            continue

        # Skip entries that look like PURE formula fragments (only math symbols)
        # But not chemical formulas like "H₂O", "CO₂"
        if re.match(r'^[=+\-*/()0-9.\s]+$', t):
            continue

        # Skip entries starting with common sentence starters that indicate fragments
        fragment_starters = [
            'The ', 'It ', 'If ', 'As ', 'In ', 'We ', 'There ', 'This ',
            'When ', 'Since ', 'But ', 'And ', 'For ', 'To ', 'From ',
            'An ', 'Its ', 'Sometimes', 'Generally', 'However', 'Moreover',
            'The changes in', 'The values of', 'The situation', 'The cost of',
            'It is', 'As per', 'n=', 'r=', 'If the', 'Atis',
        ]
        is_fragment = False
        for starter in fragment_starters:
            if t.startswith(starter) and len(t) > 60:
                # Long sentences starting with these are likely content, not topics
                is_fragment = True
                break

        if is_fragment:
            continue

        cleaned.append(t)

    return cleaned


def _clean_chapter_data(chapter):
    """Post-process a chapter dict to clean all its nested data."""
    # Clean topics in sections
    for sec in chapter.get('sections', []):
        sec['topics'] = _clean_topics(sec.get('topics', []))

        # Clean topics in sub_sections
        for subsec in sec.get('sub_sections', []):
            subsec['topics'] = _clean_topics(subsec.get('topics', []))

    # Clean exercises
    if 'exercises' in chapter:
        chapter['exercises'] = _clean_exercises(chapter['exercises'])

    # Clean learning objectives (remove the "Study X" pattern)
    if 'learning_objectives' in chapter:
        los = chapter['learning_objectives']
        if isinstance(los, list):
            cleaned_los = []
            for lo in los:
                if isinstance(lo, str) and len(lo.strip()) > 5:
                    # Enhance vague "Study X" objectives
                    lo = lo.strip()
                    if re.match(r'^Study\s+', lo, re.IGNORECASE) and len(lo) < 30:
                        # Try to make it more descriptive
                        subject = re.sub(r'^Study\s+', '', lo, flags=re.IGNORECASE).strip()
                        lo = f'Understand and apply the concepts of {subject}'
                    cleaned_los.append(lo)
            chapter['learning_objectives'] = cleaned_los

    return chapter


def _build_fallback_chapters(result):
    """Basic rule-based chapter extraction when Mistral is not available."""
    chapters = []
    chapter_re = re.compile(r'^(?:chapter|ch\.?)\s+(\d+|[ivxlc]+)[\s.:]+(.+)',
                            re.IGNORECASE)
    section_re = re.compile(r'^(\d+\.\d+(?:\.\d+)?)\s+(.+)')
    exercise_header_re = re.compile(
        r'^\s*(?:Exercise|Exercises|Practice\s+Problems|Questions)\s*(\d*[\.\d]*)\s*$',
        re.IGNORECASE
    )
    # Exercise type detection patterns
    mcq_header_re = re.compile(
        r'(?:MCQ|Multiple\s+Choice|Objective|Choose\s+the\s+correct)',
        re.IGNORECASE
    )
    fill_blank_re = re.compile(
        r'(?:Fill\s+in\s+the\s+blank|Fill\s+up|Complete\s+the\s+following)',
        re.IGNORECASE
    )
    true_false_re = re.compile(
        r'(?:True\s+or\s+False|State\s+whether\s+true|State\s+true\s+or\s+false)',
        re.IGNORECASE
    )
    match_re = re.compile(
        r'(?:Match\s+the\s+following|Match\s+the\s+columns)',
        re.IGNORECASE
    )
    one_sent_re = re.compile(
        r'(?:one\s+sentence|very\s+short|one\s+word)',
        re.IGNORECASE
    )
    short_re = re.compile(r'(?:short\s+(?:answer|question))', re.IGNORECASE)
    brief_re = re.compile(r'(?:brief|long\s+answer)', re.IGNORECASE)
    detailed_re = re.compile(r'(?:detailed|essay|explain\s+in\s+detail)', re.IGNORECASE)

    current_chapter = None
    current_text_lines = []
    current_section = None

    _empty_exercises = {
        'mcq': {'count': 0, 'questions': []},
        'assertion_reason': {'count': 0, 'questions': []},
        'fill_in_the_blanks': {'count': 0, 'questions': []},
        'true_false': {'count': 0, 'questions': []},
        'match_the_following': {'count': 0, 'questions': []},
        'very_short_answer': {'count': 0, 'questions': []},
        'short_answer': {'count': 0, 'questions': []},
        'long_answer': {'count': 0, 'questions': []},
        'numerical_problems': {'count': 0, 'questions': []},
        'diagram_based': {'count': 0, 'questions': []},
        'hots': {'count': 0, 'questions': []},
        'activity_based': {'count': 0, 'questions': []},
    }

    for p in result['pages']:
        for line in p['text'].split('\n'):
            stripped = line.strip()
            m = chapter_re.match(stripped)
            if m:
                # Save previous chapter
                if current_chapter:
                    if current_text_lines:
                        current_chapter['introduction'] = '\n'.join(
                            current_text_lines[:10]).strip()
                    chapters.append(current_chapter)
                current_chapter = {
                    'chapter_number': len(chapters) + 1,
                    'chapter_name': m.group(2).strip(),
                    'sections': [],
                    'textbook_page_start': p['page_number'],
                    'textbook_page_end': None,
                    'introduction': '',
                    'exercises': copy.deepcopy(_empty_exercises),
                    'answer_keys': {'exercise_answers': []},
                    'learning_objectives': [],
                    'chapter_summary': {
                        'key_definitions': [],
                        'important_laws': [],
                        'important_reactions': [],
                        'nature_points': [],
                        'importance_points': [],
                    },
                }
                current_text_lines = []
                current_section = None
            elif current_chapter:
                # Try to detect sections
                sec_m = section_re.match(stripped)
                if sec_m:
                    sec_num = sec_m.group(1)
                    sec_name = sec_m.group(2).strip()
                    # Determine if this is a sub-section (X.Y.Z) or section (X.Y)
                    parts = sec_num.split('.')
                    if len(parts) == 3 and current_section is not None:
                        # Sub-section - add to current section
                        current_section.setdefault('sub_sections', []).append({
                            'sub_section_number': sec_num,
                            'sub_section_name': sec_name,
                            'content': '',
                            'topics': [],
                            'formulas': [],
                            'key_terms': [],
                            'diagrams': [],
                            'chemical_equations': [],
                            'derivations': [],
                            'worked_examples': [],
                        })
                    else:
                        # New section
                        current_section = {
                            'section_number': sec_num,
                            'section_name': sec_name,
                            'content': '',
                            'sub_sections': [],
                            'topics': [],
                            'formulas': [],
                            'key_terms': [],
                            'diagrams': [],
                            'chemical_equations': [],
                            'derivations': [],
                            'worked_examples': [],
                        }
                        current_chapter['sections'].append(current_section)
                else:
                    current_text_lines.append(line)
                    current_chapter['textbook_page_end'] = p['page_number']

    if current_chapter:
        if current_text_lines:
            current_chapter['introduction'] = '\n'.join(
                current_text_lines[:10]).strip()
        chapters.append(current_chapter)

    # If no chapters detected, create a single chapter from all content
    if not chapters:
        full_text = '\n'.join(p['text'] for p in result['pages'] if p['text'])
        chapters.append({
            'chapter_number': 1,
            'chapter_name': 'Full Document',
            'sections': [],
            'textbook_page_start': 1,
            'textbook_page_end': len(result['pages']),
            'introduction': full_text[:1000],
            'exercises': copy.deepcopy(_empty_exercises),
            'answer_keys': {'exercise_answers': []},
            'learning_objectives': [],
            'chapter_summary': {
                'key_definitions': [],
                'important_laws': [],
                'important_reactions': [],
                'nature_points': [],
                'importance_points': [],
            },
        })

    return chapters


def format_textbook(result, filename, options):
    """Structure extracted text into the textbook JSON schema.
    Auto-detects whether the input is a single chapter or full book,
    and auto-detects the subject if not provided."""
    
    # Textbook metadata from user-provided options
    book_id = options.get('book_id', Path(filename).stem.lower().replace(' ', '-'))
    board = options.get('board', '')
    standard = options.get('standard', '')
    stream = options.get('stream', '')
    medium = options.get('medium', '')
    language = options.get('language', 'en')
    subject_key = options.get('subject_key', '')
    subject_name = options.get('subject_name', '')

    full_text = '\n\n'.join(
        f'--- Page {p["page_number"]} ---\n{p["text"]}'
        for p in result['pages'] if p['text']
    )
    full_text = clean_for_output(full_text)

    # === AUTO-DETECT SUBJECT if user didn't provide one ===
    if not subject_name:
        subject_name = _detect_subject(result['pages'])
        if subject_name:
            log.info(f'Subject auto-detected as: {subject_name}')
            subject_key = subject_key or subject_name.lower().replace(' ', '-')

    # === AUTO-DETECT DOCUMENT TYPE ===
    doc_type, chapter_info = _detect_document_type(full_text, result['pages'])

    # Detect exercise ranges to provide as hints
    exercise_ranges = _detect_exercise_pages(result['pages'])
    ex_ranges_str = ', '.join([f"pp.{r['start']}-{r['end']}" for r in exercise_ranges])

    # Generate subject-specific extraction hints
    subj_lower = (subject_name or '').lower()
    subject_hints = ''
    if any(s in subj_lower for s in ['physics', 'physical']):
        subject_hints = (
            'SUBJECT: PHYSICS. Pay EXTREME attention to: '
            'formulas (with units and dimensions), derivations (step-by-step), '
            'numerical solved examples, diagrams (circuit, ray, force diagrams), '
            'laws and principles, SI units. '
            'For exercises: look for numerical problems with given data, '
            'assertion-reason questions, diagram-based questions. '
        )
    elif any(s in subj_lower for s in ['chemistry', 'chemical']):
        subject_hints = (
            'SUBJECT: CHEMISTRY. Pay EXTREME attention to: '
            'chemical equations (balanced, with conditions like temp/catalyst/pressure), '
            'reaction mechanisms, IUPAC nomenclature, structural formulas, '
            'periodic table trends, electron configurations, '
            'mole calculations, stoichiometry problems. '
            'For exercises: look for balancing equations, reaction prediction, '
            'numerical problems with molar mass calculations. '
        )
    elif any(s in subj_lower for s in ['biology', 'biological', 'botany', 'zoology']):
        subject_hints = (
            'SUBJECT: BIOLOGY. Pay EXTREME attention to: '
            'classification hierarchies, process descriptions (photosynthesis, respiration, '
            'cell division, digestion), labeled diagram descriptions, '
            'difference tables, life cycles, genetic crosses, '
            'ecological concepts, anatomical structures. '
            'For exercises: look for diagram-based questions, '
            'give-reasons questions, differentiate-between questions. '
        )
    elif any(s in subj_lower for s in ['math', 'mathematics', 'algebra', 'geometry', 'calculus']):
        subject_hints = (
            'SUBJECT: MATHEMATICS. Pay EXTREME attention to: '
            'theorem statements and COMPLETE proofs, formulas and identities, '
            'worked examples with step-by-step solutions, constructions, '
            'coordinate geometry problems, trigonometric identities. '
            'For exercises: extract ALL numerical problems with given data, '
            'prove-that questions, construction problems. '
        )
    elif any(s in subj_lower for s in ['statistic']):
        subject_hints = (
            'SUBJECT: STATISTICS. Pay EXTREME attention to: '
            'index numbers, correlation, regression, time series, '
            'formulas with detailed notation, worked examples. '
        )
    else:
        subject_hints = f'Subject: {subject_name or "academic"}. '

    # Try Mistral AI structuring
    final_chapters = []
    mistral_key = options.get('mistral_api_key', '') or os.environ.get('MISTRAL_API_KEY', '')
    
    if mistral_key and MISTRAL_AVAILABLE:
        log.info(f'Using Mistral AI for textbook structuring (V3 - {doc_type} mode)...')
        try:
            from mistralai.client import Mistral as MistralClient
            client = MistralClient(api_key=mistral_key, timeout_ms=120000)
        except ImportError:
            try:
                from mistralai import Mistral as MistralClient
                client = MistralClient(api_key=mistral_key, timeout_ms=120000)
            except Exception:
                client = None

        if client:
            TARGET_CHUNK_SIZE = 40000

            if doc_type == 'single_chapter':
                # ====== SINGLE CHAPTER MODE ======
                ch_name = chapter_info['name'] if chapter_info else Path(filename).stem
                ch_num = chapter_info['number'] if chapter_info else '1'
                log.info(f'Single chapter mode: "{ch_name}" (Chapter {ch_num})')

                if len(full_text) <= TARGET_CHUNK_SIZE:
                    # Small enough - process entire chapter in one go
                    prefix = (
                        f'{subject_hints}'
                        f'This document is a SINGLE CHAPTER from a {subject_name or "textbook"}. '
                        f'Chapter: "{ch_name}" (Number: {ch_num}). '
                        f'Board: {board or "N/A"}, Standard: {standard or "N/A"}. '
                        f'There are {len(result["pages"])} pages total. '
                        f'EXERCISE HINT: Exercises likely at: {ex_ranges_str or "end of chapter"}. '
                        f'CRITICAL: This is the ENTIRE chapter. Extract EVERY section, subsection, '
                        f'formula, equation, derivation, diagram, worked example, '
                        f'and ALL exercises. The "content" field must contain the COMPLETE '
                        f'teaching material for each section. Do NOT truncate or summarize. '
                        f'Output exactly ONE chapter object with ALL its content.'
                    )
                    parsed = _call_mistral_for_structuring(client, full_text, prompt_prefix=prefix)
                    if parsed and 'chapters' in parsed:
                        final_chapters.extend(parsed['chapters'])
                    elif parsed and ('chapter_name' in parsed or 'sections' in parsed):
                        final_chapters.append(parsed)
                else:
                    # Large chapter - split into page chunks with overlap
                    page_parts = re.split(r'(--- Page \d+ ---)', full_text)
                    page_blocks = []
                    for i in range(1, len(page_parts), 2):
                        header = page_parts[i]
                        content = page_parts[i+1] if i+1 < len(page_parts) else ''
                        page_blocks.append(header + content)
                    
                    if not page_blocks:
                        page_blocks = [full_text[i:i+TARGET_CHUNK_SIZE] for i in range(0, len(full_text), TARGET_CHUNK_SIZE)]

                    total_chunks = 0
                    idx = 0
                    temp_chunks = []
                    while idx < len(page_blocks):
                        batch = []
                        batch_len = 0
                        while idx < len(page_blocks) and batch_len < TARGET_CHUNK_SIZE:
                            batch.append(page_blocks[idx])
                            batch_len += len(page_blocks[idx])
                            idx += 1
                        temp_chunks.append(batch)
                    total_chunks = len(temp_chunks)

                    chapter_parts = []
                    idx = 0
                    chunk_num = 0
                    while idx < len(page_blocks):
                        batch = []
                        batch_len = 0
                        while idx < len(page_blocks) and batch_len < TARGET_CHUNK_SIZE:
                            batch.append(page_blocks[idx])
                            batch_len += len(page_blocks[idx])
                            idx += 1
                        
                        chunk_num += 1
                        chunk_text = "\n".join(batch)
                        
                        # Add overlap
                        if idx > len(batch) and len(batch) > 0:
                            overlap_idx = idx - len(batch) - 1
                            if overlap_idx >= 0:
                                chunk_text = page_blocks[overlap_idx] + "\n[OVERLAP FROM PREVIOUS]\n" + chunk_text
                        
                        prefix = (
                            f'{subject_hints}'
                            f'This is part {chunk_num} of {total_chunks} of a SINGLE CHAPTER: '
                            f'"{ch_name}" (Chapter {ch_num}) from a {subject_name or "textbook"}. '
                            f'Board: {board or "N/A"}, Standard: {standard or "N/A"}. '
                            f'EXERCISE HINT: Exercises likely at: {ex_ranges_str or "end of chapter"}. '
                            f'Extract ALL sections, subsections, formulas, equations, derivations, '
                            f'diagrams, worked examples, and exercises found in this segment. '
                            f'The "content" field must be COMPLETE. Do not truncate.'
                        )
                        
                        parsed = _call_mistral_for_structuring(client, chunk_text, prompt_prefix=prefix)
                        if parsed and 'chapters' in parsed:
                            chapter_parts.extend(parsed['chapters'])
                        elif parsed and ('chapter_name' in parsed or 'sections' in parsed):
                            chapter_parts.append(parsed)
                    
                    _merge_chapters(final_chapters, chapter_parts)

            else:
                # ====== FULL BOOK MODE ======
                log.info('Full book mode: splitting by chapter boundaries...')
                chapter_segments = _split_text_by_chapters(full_text)
                log.info(f'Detected {len(chapter_segments)} chapter segment(s) for AI processing')

                for seg_idx, (chapter_hint, seg_text, start_page) in enumerate(chapter_segments):
                    if 'front matter' in chapter_hint.lower() and len(seg_text) < 1000:
                        continue

                    log.info(f'Processing segment {seg_idx + 1}/{len(chapter_segments)}: "{chapter_hint}"')

                    if len(seg_text) > TARGET_CHUNK_SIZE:
                        pages_split = re.split(r'(--- Page \d+ ---)', seg_text)
                        page_blocks = []
                        for i in range(1, len(pages_split), 2):
                            header = pages_split[i]
                            content = pages_split[i+1] if i+1 < len(pages_split) else ''
                            page_blocks.append(header + content)
                        
                        if not page_blocks:
                            page_blocks = [seg_text[i:i+TARGET_CHUNK_SIZE] for i in range(0, len(seg_text), TARGET_CHUNK_SIZE)]

                        chapter_parts = []
                        idx = 0
                        while idx < len(page_blocks):
                            batch = []
                            batch_len = 0
                            while idx < len(page_blocks) and batch_len < TARGET_CHUNK_SIZE:
                                batch.append(page_blocks[idx])
                                batch_len += len(page_blocks[idx])
                                idx += 1
                            
                            chunk_text = "\n".join(batch)
                            
                            if idx > len(batch) and len(batch) > 0:
                                overlap_idx = idx - len(batch) - 1
                                if overlap_idx >= 0:
                                    chunk_text = page_blocks[overlap_idx] + "\n[OVERLAP FROM PREVIOUS]\n" + chunk_text
                            
                            prefix = (
                                f'{subject_hints}'
                                f'This is a segment of chapter "{chapter_hint}" from a {subject_name or "textbook"}. '
                                f'Board: {board or "N/A"}, Standard: {standard or "N/A"}. '
                                f'EXERCISE HINT: Exercises likely at: {ex_ranges_str or "end of chapter"}. '
                                f'CRITICAL: Extract ALL content, formulas, equations, derivations, and diagrams. '
                                f'The "content" field in each section must contain the COMPLETE teaching material. '
                                f'Do not skip ANY section or subsection. Extract ALL exercises by type.'
                            )
                            
                            parsed = _call_mistral_for_structuring(client, chunk_text, prompt_prefix=prefix)
                            if parsed and 'chapters' in parsed:
                                chapter_parts.extend(parsed['chapters'])
                            elif parsed and ('chapter_name' in parsed or 'sections' in parsed):
                                chapter_parts.append(parsed)
                        
                        _merge_chapters(final_chapters, chapter_parts)
                    else:
                        prefix = (
                            f'{subject_hints}'
                            f'This is a COMPLETE chapter: "{chapter_hint}". '
                            f'Board: {board or "N/A"}, Standard: {standard or "N/A"}. '
                            f'EXERCISE HINT: Exercises likely at: {ex_ranges_str or "end of chapter"}. '
                            f'Extract EVERY section, subsection, formula, equation, derivation, diagram, '
                            f'worked example, and ALL exercises by type. No truncation allowed. '
                            f'The "content" field must contain the COMPLETE teaching material for each section.'
                        )
                        parsed = _call_mistral_for_structuring(client, seg_text, prompt_prefix=prefix)
                        if parsed and 'chapters' in parsed:
                            _merge_chapters(final_chapters, parsed['chapters'])
                        elif parsed and ('chapter_name' in parsed or 'sections' in parsed):
                            _merge_chapters(final_chapters, [parsed])

    # Fallback if no chapters detected
    if not final_chapters:
        log.info('Using rule-based textbook structuring (AI failed or was missing)...')
        final_chapters = _build_fallback_chapters(result)

    # Post-process all chapters: clean garbage data
    for ch in final_chapters:
        ch = _clean_chapter_data(ch)

    # Ensure all chapters have the required structure
    _empty_exercises = {
        'mcq': {'count': 0, 'questions': []},
        'assertion_reason': {'count': 0, 'questions': []},
        'fill_in_the_blanks': {'count': 0, 'questions': []},
        'true_false': {'count': 0, 'questions': []},
        'match_the_following': {'count': 0, 'questions': []},
        'very_short_answer': {'count': 0, 'questions': []},
        'short_answer': {'count': 0, 'questions': []},
        'long_answer': {'count': 0, 'questions': []},
        'numerical_problems': {'count': 0, 'questions': []},
        'diagram_based': {'count': 0, 'questions': []},
        'hots': {'count': 0, 'questions': []},
        'activity_based': {'count': 0, 'questions': []},
    }

    # New section-level fields for science subjects
    _section_fields = {
        'sub_sections': [], 'topics': [], 'content': '', 'formulas': [],
        'key_terms': [], 'diagrams': [], 'chemical_equations': [],
        'derivations': [], 'worked_examples': [],
    }

    for ch in final_chapters:
        ch.setdefault('sections', [])
        ch.setdefault('textbook_page_start', None)
        ch.setdefault('textbook_page_end', None)
        ch.setdefault('introduction', '')
        
        existing_ex = ch.get('exercises', {})
        if not isinstance(existing_ex, dict):
            existing_ex = {}
        for ex_type, ex_default in _empty_exercises.items():
            existing_ex.setdefault(ex_type, copy.deepcopy(ex_default))
        ch['exercises'] = existing_ex
        
        ch.setdefault('answer_keys', {'exercise_answers': []})
        ch.setdefault('learning_objectives', [])
        ch.setdefault('chapter_summary', {
            'key_definitions': [],
            'important_laws': [],
            'important_reactions': [],
            'nature_points': [],
            'importance_points': [],
        })
        
        for sec in ch.get('sections', []):
            for field, default in _section_fields.items():
                sec.setdefault(field, copy.deepcopy(default) if isinstance(default, list) else default)
            for subsec in sec.get('sub_sections', []):
                for field, default in _section_fields.items():
                    if field != 'sub_sections':
                        subsec.setdefault(field, copy.deepcopy(default) if isinstance(default, list) else default)

    return {
        'book_id': book_id,
        'board': board,
        'standard': standard,
        'stream': stream,
        'medium': medium,
        'language': language,
        'subject_key': subject_key,
        'subject_name': subject_name,
        'total_chapters': len(final_chapters),
        'meta': {
            'publisher': options.get('publisher', ''),
            'additional_sections': [],
            'authors': [],
            'subject_adviser': '',
            'additional_content': {},
            'llm_context_notes': {
                'purpose': 'This JSON provides structured textbook context for LLM consumption.',
                'subject_overview': f'{subject_name} is a {stream} stream subject for Standard {standard}.',
                'exam_pattern': '',
                'key_management_thinkers': [],
                'important_acts_and_laws': [],
                'important_organisations': [],
            },
        },
        'chapters': final_chapters,
    }


FORMAT_HANDLERS = {
    'raw': format_raw,
    'chunks': format_chunks,
    'paragraphs': format_paragraphs,
    'sentences': format_sentences,
    'jsonl': format_jsonl,
    'conversation': format_conversation,
    'textbook': format_textbook,
}


# =====================================================================
# AI VALIDATOR HELPERS (GROQ)
# =====================================================================

PATH_TOKEN_RE = re.compile(r'([A-Za-z0-9_\-]+)|\[(\d+)\]')


def _strip_code_fences(text):
    cleaned = (text or '').strip()
    if cleaned.startswith('```'):
        cleaned = re.sub(r'^```(?:json)?', '', cleaned, flags=re.IGNORECASE).strip()
        if cleaned.endswith('```'):
            cleaned = cleaned[:-3].strip()
    return cleaned


def _safe_json_loads(text):
    cleaned = _strip_code_fences(text)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        first = cleaned.find('{')
        last = cleaned.rfind('}')
        if first != -1 and last != -1 and last > first:
            return json.loads(cleaned[first:last + 1])
        raise


def _to_score(value, default_value=0.0):
    try:
        num = float(value)
        return round(max(1.0, min(10.0, num)), 2)
    except Exception:
        return default_value


def _normalize_suggestions(suggestions):
    normalized = []
    if not isinstance(suggestions, list):
        return normalized

    for idx, suggestion in enumerate(suggestions[:8], start=1):
        if not isinstance(suggestion, dict):
            continue

        sid = str(suggestion.get('id') or f'S{idx}')[:40]
        title = str(suggestion.get('title') or f'Suggestion {idx}')[:180]
        reason = str(suggestion.get('reason') or '').strip()[:600]
        impact = str(suggestion.get('expected_impact') or '').strip()[:300]

        raw_ops = suggestion.get('operations', [])
        operations = []
        if isinstance(raw_ops, list):
            for op in raw_ops[:3]:
                if not isinstance(op, dict):
                    continue
                op_name = str(op.get('op') or '').strip().lower()
                path = str(op.get('path') or '').strip()
                if not path:
                    continue

                if op_name in ('set', 'add'):
                    if 'value' not in op:
                        continue
                    operations.append({
                        'op': 'set',
                        'path': path,
                        'value': op.get('value'),
                    })
                elif op_name == 'replace_substring':
                    search = str(op.get('search') or '')
                    if not search:
                        continue
                    operations.append({
                        'op': 'replace_substring',
                        'path': path,
                        'search': search,
                        'replace': str(op.get('replace') or ''),
                    })
                elif op_name == 'append_list':
                    if 'value' not in op:
                        continue
                    operations.append({
                        'op': 'append_list',
                        'path': path,
                        'value': op.get('value'),
                    })
                elif op_name == 'remove':
                    operations.append({
                        'op': 'remove',
                        'path': path,
                    })

        normalized.append({
            'id': sid,
            'title': title,
            'reason': reason,
            'expected_impact': impact,
            'operations': operations,
        })

    return normalized


def _normalize_validator_response(raw_data):
    if not isinstance(raw_data, dict):
        raw_data = {}

    score = _to_score(raw_data.get('score'), 0.0)
    structure_score = _to_score(raw_data.get('structure_score'), score)
    accuracy_score = _to_score(raw_data.get('accuracy_score'), score)
    reliability_score = _to_score(raw_data.get('reliability_score'), score)

    issues = raw_data.get('issues', [])
    if not isinstance(issues, list):
        issues = []
    issues = [str(x)[:260] for x in issues[:10] if str(x).strip()]

    return {
        'score': score,
        'structure_score': structure_score,
        'accuracy_score': accuracy_score,
        'reliability_score': reliability_score,
        'summary': str(raw_data.get('summary') or '').strip()[:1000],
        'issues': issues,
        'suggestions': _normalize_suggestions(raw_data.get('suggestions')),
    }


def _collect_text_fragments(node, out, max_fragments=2500):
    if len(out) >= max_fragments:
        return

    if isinstance(node, dict):
        for key, value in node.items():
            if key in {'id', 'checksum', 'created_at', 'source_file', 'format'}:
                continue
            _collect_text_fragments(value, out, max_fragments)
            if len(out) >= max_fragments:
                return
    elif isinstance(node, list):
        for item in node:
            _collect_text_fragments(item, out, max_fragments)
            if len(out) >= max_fragments:
                return
    elif isinstance(node, str):
        txt = ' '.join(node.split())
        if txt and any(ch.isalpha() for ch in txt):
            out.append(txt[:1200])


def _extract_text_for_validation(payload, max_chars=160000):
    fragments = []
    _collect_text_fragments(payload, fragments)

    joined = []
    total = 0
    for frag in fragments:
        if total >= max_chars:
            break
        remaining = max_chars - total
        part = frag[:remaining]
        joined.append(part)
        total += len(part) + 1
    return '\n'.join(joined)


def _load_output_payload(output_filename):
    safe_name = Path(output_filename or '').name
    if not safe_name:
        raise ValueError('Missing output file name')

    output_path = os.path.join(app.config['OUTPUT_FOLDER'], safe_name)
    if not os.path.exists(output_path):
        raise FileNotFoundError(f'Output file not found: {safe_name}')

    is_jsonl = safe_name.endswith('.jsonl')
    if is_jsonl:
        lines = []
        with open(output_path, 'r', encoding='utf-8') as f:
            for line_no, line in enumerate(f, start=1):
                stripped = line.strip()
                if not stripped:
                    continue
                try:
                    lines.append(json.loads(stripped))
                except json.JSONDecodeError as e:
                    raise ValueError(f'Invalid JSONL at line {line_no}: {e.msg}') from e
        payload = {
            'format': 'jsonl',
            'lines': lines,
        }
    else:
        with open(output_path, 'r', encoding='utf-8') as f:
            payload = json.load(f)

    return safe_name, output_path, payload, is_jsonl


def _payload_preview(payload, is_jsonl):
    if is_jsonl:
        lines = payload.get('lines', []) if isinstance(payload, dict) else []
        text = '\n'.join(json.dumps(x, ensure_ascii=False) for x in lines)
        return text[:5000]

    text = json.dumps(payload, indent=2, ensure_ascii=False)
    if len(text) > 50000:
        return text[:50000] + '\n\n... [TRUNCATED - download full file] ...'
    return text


def _serialize_payload(payload, is_jsonl):
    if is_jsonl:
        if not isinstance(payload, dict) or not isinstance(payload.get('lines'), list):
            raise ValueError('Invalid JSONL payload structure')
        return '\n'.join(json.dumps(x, ensure_ascii=False) for x in payload['lines'])
    return json.dumps(payload, indent=2, ensure_ascii=False)


def _tokenize_path(path):
    cleaned = (path or '').strip()
    if cleaned.startswith('$.'):
        cleaned = cleaned[2:]
    elif cleaned.startswith('$'):
        cleaned = cleaned[1:]
    if not cleaned:
        return []

    tokens = []
    for key, idx in PATH_TOKEN_RE.findall(cleaned):
        if key:
            tokens.append(key)
        elif idx:
            tokens.append(int(idx))
    return tokens


def _get_value_at_path(data, path):
    tokens = _tokenize_path(path)
    cur = data
    for token in tokens:
        if isinstance(token, int):
            if not isinstance(cur, list) or token < 0 or token >= len(cur):
                raise KeyError(f'Invalid list index at path: {path}')
            cur = cur[token]
        else:
            if not isinstance(cur, dict) or token not in cur:
                raise KeyError(f'Missing key at path: {path}')
            cur = cur[token]
    return cur


def _get_parent_and_leaf(data, path):
    tokens = _tokenize_path(path)
    if not tokens:
        raise ValueError('Path cannot be root')

    parent = data
    for token in tokens[:-1]:
        if isinstance(token, int):
            if not isinstance(parent, list) or token < 0 or token >= len(parent):
                raise KeyError(f'Invalid list index at path: {path}')
            parent = parent[token]
        else:
            if not isinstance(parent, dict) or token not in parent:
                raise KeyError(f'Missing key at path: {path}')
            parent = parent[token]
    return parent, tokens[-1]


def _set_value_at_path(data, path, value):
    parent, leaf = _get_parent_and_leaf(data, path)
    if isinstance(leaf, int):
        if not isinstance(parent, list) or leaf < 0 or leaf >= len(parent):
            raise KeyError(f'Invalid list index at path: {path}')
        parent[leaf] = value
    else:
        if not isinstance(parent, dict):
            raise KeyError(f'Parent is not an object at path: {path}')
        parent[leaf] = value


def _apply_operation(payload, operation):
    if not isinstance(operation, dict):
        raise ValueError('Operation must be a JSON object')

    op = str(operation.get('op') or '').strip().lower()
    path = str(operation.get('path') or '').strip()
    if not path:
        raise ValueError('Operation path is required')

    if op == 'set':
        if 'value' not in operation:
            raise ValueError('set operation requires value')
        _set_value_at_path(payload, path, operation.get('value'))
        return

    if op == 'replace_substring':
        search = str(operation.get('search') or '')
        replace_with = str(operation.get('replace') or '')
        if not search:
            raise ValueError('replace_substring requires search')
        current = _get_value_at_path(payload, path)
        if not isinstance(current, str):
            raise ValueError(f'Path is not a string: {path}')
        if search not in current:
            raise ValueError('search text not found at path')
        updated = current.replace(search, replace_with, 1)
        _set_value_at_path(payload, path, updated)
        return

    if op == 'append_list':
        if 'value' not in operation:
            raise ValueError('append_list requires value')
        target = _get_value_at_path(payload, path)
        if not isinstance(target, list):
            raise ValueError(f'Path is not a list: {path}')
        target.append(operation.get('value'))
        return

    if op == 'remove':
        parent, leaf = _get_parent_and_leaf(payload, path)
        if isinstance(leaf, int):
            if not isinstance(parent, list) or leaf < 0 or leaf >= len(parent):
                raise KeyError(f'Invalid list index at path: {path}')
            parent.pop(leaf)
        else:
            if not isinstance(parent, dict) or leaf not in parent:
                raise KeyError(f'Missing key at path: {path}')
            parent.pop(leaf)
        return

    raise ValueError(f'Unsupported operation: {op}')


def _apply_suggestions(payload, suggestions):
    applied = []
    skipped = []

    if not isinstance(suggestions, list):
        return applied, skipped

    for idx, suggestion in enumerate(suggestions, start=1):
        if not isinstance(suggestion, dict):
            skipped.append({
                'id': f'S{idx}',
                'reason': 'Suggestion is not an object',
            })
            continue

        sid = str(suggestion.get('id') or f'S{idx}')
        title = str(suggestion.get('title') or sid)
        operations = suggestion.get('operations', [])
        if not isinstance(operations, list) or not operations:
            skipped.append({
                'id': sid,
                'title': title,
                'reason': 'No operations provided',
            })
            continue

        applied_ops = 0
        op_errors = []
        for op in operations:
            try:
                _apply_operation(payload, op)
                applied_ops += 1
            except Exception as exc:
                op_errors.append(str(exc))

        if applied_ops > 0:
            applied.append({
                'id': sid,
                'title': title,
                'operations_applied': applied_ops,
                'operation_errors': op_errors,
            })
        else:
            skipped.append({
                'id': sid,
                'title': title,
                'reason': '; '.join(op_errors) if op_errors else 'No operations could be applied',
            })

    return applied, skipped


def _build_validator_user_prompt(output_file, payload):
    extracted_text = _extract_text_for_validation(payload)
    if not extracted_text.strip():
        raise ValueError('No extractable text found to validate')

    json_sample = json.dumps(payload, ensure_ascii=False)
    if len(json_sample) > 50000:
        json_sample = json_sample[:50000] + '\n... [TRUNCATED] ...'

    text_sample = extracted_text[:160000]
    return (
        f'Validate this converted output file: {output_file}.\n\n'
        'Scoring requirements:\n'
        '1) Structure consistency and schema alignment\n'
        '2) Text accuracy and OCR cleanliness\n'
        '3) Reliability for downstream AI training\n\n'
        'Important constraints:\n'
        '- Validation must be read-only.\n'
        '- Suggestions should improve score toward 10/10.\n'
        '- Suggestions must be patch-like and safe.\n'
        '- If source is JSONL, root contains key "lines".\n\n'
        f'Extra note: {MUTUAL_DECISION_NOTE}\n\n'
        f'JSON sample:\n{json_sample}\n\n'
        f'Extracted text sample:\n{text_sample}'
    )


def _run_provider_validator(provider, api_key, base_url, models, output_file, payload):
    if not OPENAI_AVAILABLE:
        raise RuntimeError('OpenAI Python package is required for provider API calls')
    if not api_key:
        raise ValueError(f'{provider} API key is required')

    user_prompt = _build_validator_user_prompt(output_file, payload)
    client = OpenAI(api_key=api_key, base_url=base_url)
    last_error = None

    for model in models:
        try:
            try:
                response = client.chat.completions.create(
                    model=model,
                    temperature=0.0,
                    response_format={'type': 'json_object'},
                    messages=[
                        {'role': 'system', 'content': VALIDATOR_SYSTEM_PROMPT},
                        {'role': 'user', 'content': user_prompt},
                    ],
                )
            except Exception:
                response = client.chat.completions.create(
                    model=model,
                    temperature=0.0,
                    messages=[
                        {'role': 'system', 'content': VALIDATOR_SYSTEM_PROMPT},
                        {'role': 'user', 'content': user_prompt},
                    ],
                )

            content = response.choices[0].message.content if response.choices else ''
            parsed = _safe_json_loads(content)
            normalized = _normalize_validator_response(parsed)
            normalized['provider'] = provider
            normalized['model'] = model
            return normalized
        except Exception as exc:
            last_error = exc
            log.warning(f'{provider} validator call failed for model {model}: {exc}')

    raise RuntimeError(f'{provider} validator failed for all models: {last_error}')


def _run_groq_validator(groq_api_key, output_file, payload):
    return _run_provider_validator(
        provider='groq',
        api_key=groq_api_key,
        base_url=GROQ_BASE_URL,
        models=GROQ_MODELS,
        output_file=output_file,
        payload=payload,
    )


def _run_sarvam_validator(sarvam_api_key, output_file, payload):
    return _run_provider_validator(
        provider='sarvam',
        api_key=sarvam_api_key,
        base_url=SARVAM_BASE_URL,
        models=SARVAM_MODELS,
        output_file=output_file,
        payload=payload,
    )


def _text_key(value):
    return re.sub(r'\s+', ' ', re.sub(r'[^a-z0-9]+', ' ', str(value).lower())).strip()


def _merge_issues_mutually(primary, secondary):
    buckets = {}

    for source, issues in (('groq', primary or []), ('sarvam', secondary or [])):
        for idx, issue in enumerate(issues):
            text = str(issue).strip()
            key = _text_key(text)
            if not key:
                continue

            row = buckets.setdefault(key, {
                'text': text,
                'sources': set(),
                'rank': 0,
            })
            if len(text) < len(row['text']):
                row['text'] = text
            row['sources'].add(source)
            row['rank'] += max(1, 10 - idx)

    ranked = sorted(
        buckets.values(),
        key=lambda row: (len(row['sources']), row['rank']),
        reverse=True,
    )
    return [row['text'] for row in ranked[:10]]


def _suggestion_signature(suggestion):
    if not isinstance(suggestion, dict):
        return ''

    operations = suggestion.get('operations', [])
    op_bits = []
    if isinstance(operations, list):
        for op in operations:
            if not isinstance(op, dict):
                continue
            op_bits.append(f"{op.get('op', '')}:{op.get('path', '')}")
    ops_key = '|'.join(op_bits)

    title_key = _text_key(suggestion.get('title', ''))
    reason_key = _text_key(suggestion.get('reason', ''))[:120]
    return f'{ops_key}::{title_key}::{reason_key}'


def _merge_suggestions_mutually(groq_suggestions, sarvam_suggestions, groq_score, sarvam_score):
    grouped = {}

    for provider, items in (('groq', groq_suggestions or []), ('sarvam', sarvam_suggestions or [])):
        if not isinstance(items, list):
            continue
        for suggestion in items:
            if not isinstance(suggestion, dict):
                continue
            sig = _suggestion_signature(suggestion)
            if not sig:
                continue

            row = grouped.setdefault(sig, {
                'sources': set(),
                'items': [],
            })
            row['sources'].add(provider)
            row['items'].append((provider, suggestion))

    weights = {'groq': groq_score, 'sarvam': sarvam_score}
    merged = []

    for idx, row in enumerate(grouped.values(), start=1):
        picked_provider, picked = max(
            row['items'],
            key=lambda item: (
                weights.get(item[0], 0),
                len(item[1].get('operations', []) if isinstance(item[1].get('operations', []), list) else []),
            ),
        )

        suggestion = copy.deepcopy(picked)
        suggestion['id'] = str(suggestion.get('id') or f'M{idx}')
        suggestion['agreed_by'] = sorted(row['sources'])
        if len(row['sources']) > 1:
            impact = str(suggestion.get('expected_impact') or 'High impact').strip()
            suggestion['expected_impact'] = f'{impact} (mutual agreement)'
        else:
            suggestion['expected_impact'] = str(
                suggestion.get('expected_impact') or f'Provider preference: {picked_provider}'
            )
        merged.append(suggestion)

    merged.sort(
        key=lambda s: (
            len(s.get('agreed_by', [])),
            len(s.get('operations', []) if isinstance(s.get('operations', []), list) else []),
        ),
        reverse=True,
    )
    return merged[:8]


def _build_mutual_validator(groq_result, sarvam_result):
    groq_score = _to_score(groq_result.get('score'), 0.0)
    sarvam_score = _to_score(sarvam_result.get('score'), 0.0)

    combined_score = round((groq_score + sarvam_score) / 2, 2)
    structure_score = round((
        _to_score(groq_result.get('structure_score'), groq_score)
        + _to_score(sarvam_result.get('structure_score'), sarvam_score)
    ) / 2, 2)
    accuracy_score = round((
        _to_score(groq_result.get('accuracy_score'), groq_score)
        + _to_score(sarvam_result.get('accuracy_score'), sarvam_score)
    ) / 2, 2)
    reliability_score = round((
        _to_score(groq_result.get('reliability_score'), groq_score)
        + _to_score(sarvam_result.get('reliability_score'), sarvam_score)
    ) / 2, 2)

    disagreement = round(abs(groq_score - sarvam_score), 2)
    confidence = 'high' if disagreement <= 1.0 else ('medium' if disagreement <= 2.0 else 'low')

    issues = _merge_issues_mutually(
        groq_result.get('issues', []),
        sarvam_result.get('issues', []),
    )
    suggestions = _merge_suggestions_mutually(
        groq_result.get('suggestions', []),
        sarvam_result.get('suggestions', []),
        groq_score=groq_score,
        sarvam_score=sarvam_score,
    )

    return {
        'score': combined_score,
        'structure_score': structure_score,
        'accuracy_score': accuracy_score,
        'reliability_score': reliability_score,
        'summary': (
            f'Mutual decision: Groq scored {groq_score:.1f}/10 and Sarvam scored {sarvam_score:.1f}/10. '
            f'Consensus confidence is {confidence}. '
            f'Combined score is {combined_score:.1f}/10.'
        ),
        'issues': issues,
        'suggestions': suggestions,
        'provider': 'mutual',
        'model': 'groq+sarvam-consensus',
        'consensus': {
            'disagreement': disagreement,
            'confidence': confidence,
        },
    }


# =====================================================================
# ROUTES
# =====================================================================

@app.route('/')
def index():
    return render_template('index.html', ocr_available=OCR_AVAILABLE)


@app.route('/api/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({
            'error': f'Invalid file type. Allowed: {", ".join(ALLOWED_EXTENSIONS)}'
        }), 400

    # Options
    output_format = request.form.get('format', 'chunks')
    chunk_size = int(request.form.get('chunk_size', 512))
    chunk_overlap = int(request.form.get('chunk_overlap', 64))
    system_prompt = request.form.get('system_prompt', '')
    extract_images = request.form.get('extract_images', 'false') == 'true'
    include_base64 = request.form.get('include_base64', 'false') == 'true'
    mistral_api_key = request.form.get('mistral_api_key', '')

    options = {
        'chunk_size': chunk_size,
        'chunk_overlap': chunk_overlap,
        'extract_images': extract_images,
        'include_base64': include_base64,
        'mistral_api_key': mistral_api_key,
    }
    if system_prompt:
        options['system_prompt'] = system_prompt

    # Textbook metadata fields
    if output_format == 'textbook':
        options['book_id'] = request.form.get('book_id', '')
        options['board'] = request.form.get('board', '')
        options['standard'] = request.form.get('standard', '')
        options['stream'] = request.form.get('stream', '')
        options['medium'] = request.form.get('medium', '')
        options['language'] = request.form.get('language', 'en')
        options['subject_key'] = request.form.get('subject_key', '')
        options['subject_name'] = request.form.get('subject_name', '')
        options['publisher'] = request.form.get('publisher', '')

    # Save uploaded file
    filename = file.filename
    safe_name = Path(filename).name
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_name)
    file.save(filepath)

    try:
        # Run the extraction pipeline
        processor = DocumentProcessor(filepath, safe_name, options)
        result = processor.process()

        if not result['pages'] or all(
                not p['text'] for p in result['pages']):
            return jsonify({
                'error': 'Could not extract any text from this file.'
            }), 422

        # Format output
        handler = FORMAT_HANDLERS.get(output_format, format_chunks)
        formatted = handler(result, safe_name, options)

        # Determine output filename
        base_name = Path(safe_name).stem
        is_jsonl = output_format == 'jsonl'
        ext = 'jsonl' if is_jsonl else 'json'
        output_filename = f'{base_name}_{output_format}.{ext}'
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

        # Write to disk
        with open(output_path, 'w', encoding='utf-8') as f:
            if is_jsonl:
                f.write(formatted)
            else:
                json.dump(formatted, f, indent=2, ensure_ascii=False)

        # Build preview (truncated for very large books)
        if is_jsonl:
            preview = formatted[:5000]
        else:
            # For large output, truncate preview to avoid huge responses
            preview_json = json.dumps(formatted, indent=2, ensure_ascii=False)
            if len(preview_json) > 50000:
                preview = preview_json[:50000] + '\n\n... [TRUNCATED - download full file] ...'
            else:
                preview = preview_json

        # Stats
        quality = result['quality']
        stats = {
            'total_words': quality.get('total_words', 0),
            'total_pages': quality.get('total_pages', 0),
            'total_characters': quality.get('total_characters', 0),
            'quality_score': quality.get('overall_score', 0),
            'extraction_passes': quality.get('extraction_passes', 1),
            'chapters_detected': quality.get('chapters_detected', 0),
            'warnings': quality.get('warnings', []),
        }
        if output_format == 'textbook':
            stats['total_chapters'] = formatted.get('total_chapters', 0)
        elif output_format == 'chunks':
            stats['total_chunks'] = formatted.get('total_chunks', 0)
        elif output_format == 'paragraphs':
            stats['total_paragraphs'] = formatted.get('total_paragraphs', 0)
        elif output_format == 'sentences':
            stats['total_sentences'] = formatted.get('total_sentences', 0)
        elif output_format == 'conversation':
            stats['total_conversations'] = formatted.get(
                'total_conversations', 0)
        elif is_jsonl:
            stats['total_lines'] = len(formatted.strip().split('\n'))

        output_size = os.path.getsize(output_path)

        return jsonify({
            'success': True,
            'output_file': output_filename,
            'output_format': output_format,
            'output_size': output_size,
            'output_size_human': _fmt_bytes(output_size),
            'stats': stats,
            'validator_available': OPENAI_AVAILABLE,
            'preview': preview,
        })

    except Exception as e:
        log.exception('Conversion error')
        return jsonify({'error': str(e)}), 500
    finally:
        if os.path.exists(filepath):
            os.remove(filepath)


@app.route('/api/validator/score', methods=['POST'])
def validator_score():
    body = request.get_json(silent=True) or {}
    output_file = body.get('output_file', '')
    groq_api_key = (body.get('groq_api_key') or os.environ.get('GROQ_API_KEY', '')).strip()
    sarvam_api_key = (body.get('sarvam_api_key') or os.environ.get('SARVAM_API_KEY', '')).strip()

    if not output_file:
        return jsonify({'error': 'output_file is required'}), 400
    if not groq_api_key and not sarvam_api_key:
        return jsonify({'error': 'Provide Groq and/or Sarvam API key for AI validation'}), 400

    try:
        safe_name, _, payload, is_jsonl = _load_output_payload(output_file)

        provider_results = {}
        provider_errors = {}
        futures = {}

        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            if groq_api_key:
                futures[executor.submit(_run_groq_validator, groq_api_key, safe_name, payload)] = 'groq'
            if sarvam_api_key:
                futures[executor.submit(_run_sarvam_validator, sarvam_api_key, safe_name, payload)] = 'sarvam'

            for future in concurrent.futures.as_completed(futures):
                provider = futures[future]
                try:
                    provider_results[provider] = future.result()
                except Exception as exc:
                    provider_errors[provider] = str(exc)
                    log.warning(f'{provider} validator execution failed: {exc}')

        if not provider_results:
            return jsonify({
                'error': 'All validator providers failed',
                'provider_errors': provider_errors,
            }), 502

        if 'groq' in provider_results and 'sarvam' in provider_results:
            validation = _build_mutual_validator(
                groq_result=provider_results['groq'],
                sarvam_result=provider_results['sarvam'],
            )
            decision_mode = 'mutual'
        else:
            only_provider = next(iter(provider_results.keys()))
            validation = provider_results[only_provider]
            decision_mode = 'single'

        return jsonify({
            'success': True,
            'validated_file': safe_name,
            'mode': 'jsonl' if is_jsonl else 'json',
            'decision_mode': decision_mode,
            'validator': validation,
            'individual_validators': provider_results,
            'provider_errors': provider_errors,
            'suggestion_count': len(validation.get('suggestions', [])),
            'note': 'Validation is read-only. No file changes were made.',
        })
    except FileNotFoundError as e:
        return jsonify({'error': str(e)}), 404
    except Exception as e:
        log.exception('Validator score failed')
        return jsonify({'error': str(e)}), 500


@app.route('/api/validator/apply', methods=['POST'])
def validator_apply():
    body = request.get_json(silent=True) or {}
    output_file = body.get('output_file', '')
    suggestions = body.get('suggestions', [])

    if not output_file:
        return jsonify({'error': 'output_file is required'}), 400
    if not isinstance(suggestions, list) or not suggestions:
        return jsonify({'error': 'At least one suggestion is required'}), 400

    try:
        safe_name, _, payload, is_jsonl = _load_output_payload(output_file)
        updated_payload = copy.deepcopy(payload)

        applied, skipped = _apply_suggestions(updated_payload, suggestions)
        if not applied:
            return jsonify({
                'error': 'No suggestions could be applied',
                'applied': applied,
                'skipped': skipped,
            }), 422

        original = Path(safe_name)
        suffix = original.suffix
        ts = datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')
        new_filename = f'{original.stem}_validated_{ts}{suffix}'
        new_path = os.path.join(app.config['OUTPUT_FOLDER'], new_filename)

        serialized = _serialize_payload(updated_payload, is_jsonl)
        with open(new_path, 'w', encoding='utf-8') as f:
            f.write(serialized)

        output_size = os.path.getsize(new_path)

        return jsonify({
            'success': True,
            'source_output_file': safe_name,
            'output_file': new_filename,
            'output_size': output_size,
            'output_size_human': _fmt_bytes(output_size),
            'preview': _payload_preview(updated_payload, is_jsonl),
            'applied': applied,
            'skipped': skipped,
            'applied_count': len(applied),
            'skipped_count': len(skipped),
            'note': 'Applied output was saved as a new file. Original file is unchanged.',
        })
    except FileNotFoundError as e:
        return jsonify({'error': str(e)}), 404
    except Exception as e:
        log.exception('Validator apply failed')
        return jsonify({'error': str(e)}), 500


@app.route('/api/download/<filename>')
def download(filename):
    """Download a converted JSON file."""
    safe_name = Path(filename).name  # prevent directory traversal
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], safe_name)
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    if safe_name.endswith('.jsonl'):
        mimetype = 'application/x-ndjson'
    else:
        mimetype = 'application/json'

    return send_file(
        filepath,
        as_attachment=True,
        download_name=safe_name,
        mimetype=mimetype,
    )


@app.route('/api/status')
def status():
    """System status check."""
    return jsonify({
        'mistral_sdk': MISTRAL_AVAILABLE,
        'tesseract': TESSERACT_AVAILABLE,
        'ocr_available': OCR_AVAILABLE,
        'ai_validator_available': OPENAI_AVAILABLE,
        'ai_validator_provider': 'groq+sarvam',
        'engines': ['pymupdf', 'pdfplumber', 'mistral-ocr'],
        'formats': list(FORMAT_HANDLERS.keys()),
        'max_upload_mb': app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024),
    })


def _fmt_bytes(b):
    if b == 0:
        return '0 B'
    k = 1024
    sizes = ['B', 'KB', 'MB', 'GB']
    i = int(math.floor(math.log(b, k)))
    return f'{b / k ** i:.1f} {sizes[i]}'


# =====================================================================
# MAIN
# =====================================================================

if __name__ == '__main__':
    print('')
    print('  +------------------------------------------------+')
    print('  |  Doc2JSON v2 -- Advanced AI Data Converter      |')
    print('  |  http://localhost:5000                           |')
    print('  |                                                  |')
    if MISTRAL_AVAILABLE:
        print('  |  [OK] Mistral AI SDK ready                      |')
    else:
        print('  |  [!!] Mistral SDK not found (pip install it)    |')
    if TESSERACT_AVAILABLE:
        print('  |  [OK] Tesseract OCR detected (fallback)         |')
    else:
        print('  |  [--] Tesseract not found (optional fallback)   |')
    print('  +------------------------------------------------+')
    print('')
    app.run(debug=True, port=5000)
